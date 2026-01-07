# -*- coding: utf-8 -*-
"""修复专利01的段落编号"""

from docx import Document
import re

doc_path = 'D:/code/ai-bookkeeping/docs/patents/专利01_FIFO资源池钱龄计算方法_v1.2.docx'

doc = Document(doc_path)

# 记录需要重新编号的段落
# 由于删除了[0011]，从[0012]开始的段落需要减1

renumber_map = {}
for old_num in range(12, 120):  # 假设最大到120
    renumber_map[f'[{old_num:04d}]'] = f'[{old_num-1:04d}]'

modified_count = 0
for para in doc.paragraphs:
    text = para.text
    for old, new in renumber_map.items():
        if text.startswith(old):
            new_text = text.replace(old, new, 1)
            para.clear()
            para.add_run(new_text)
            print(f'重编号: {old} -> {new}')
            modified_count += 1
            break

print(f'\n共重编号 {modified_count} 个段落')

# 保存
doc.save(doc_path)
print(f'已保存: {doc_path}')
