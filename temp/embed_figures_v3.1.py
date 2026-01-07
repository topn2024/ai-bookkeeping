# -*- coding: utf-8 -*-
"""将附图嵌入到专利01 v3.1文档中"""

from docx import Document
from docx.shared import Inches
import os

doc_path = 'D:/code/ai-bookkeeping/docs/patents/专利01_FIFO资源池钱龄计算方法_v3.1.docx'
figures_dir = 'D:/code/ai-bookkeeping/docs/patents/figures/patent_01_v3.1'

# 打开文档
doc = Document(doc_path)

# 图片列表
figures = [
    '图1_FIFO资源池模型架构示意图.png',
    '图2_FIFO消耗算法流程图.png',
    '图3_增量计算优化原理图.png',
    '图4_系统分层架构图.png',
    '图5_消费链路桑基图示例.png',
    '图6_性能对比测试结果图.png',
    '图7_钱龄健康等级映射图.png',
    '图8_数据库ER图.png',
    '图9_分片存储架构图.png',
    '图10_GPU加速并行计算示意图.png',
    '图11_FPGA流水线架构图.png',
    '图12_FIFO变体应用场景图.png',
]

# 在文档末尾添加附图
doc.add_heading('附图', level=1)

for fig_name in figures:
    fig_path = os.path.join(figures_dir, fig_name)
    if os.path.exists(fig_path):
        # 添加图片标题
        title = fig_name.replace('.png', '').replace('_', ' ')
        doc.add_paragraph(title, style='Caption')
        # 添加图片
        doc.add_picture(fig_path, width=Inches(6))
        doc.add_paragraph('')  # 空行
        print(f'  已嵌入: {fig_name}')
    else:
        print(f'  未找到: {fig_name}')

# 保存
doc.save(doc_path)
print(f'\n附图已嵌入到: {doc_path}')
