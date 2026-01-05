# -*- coding: utf-8 -*-
"""
将生成的附图嵌入到专利文档中
专利06-12使用子目录附图
"""

import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

PATENTS_DIR = 'D:/code/ai-bookkeeping/docs/patents'
FIGURES_DIR = 'D:/code/ai-bookkeeping/docs/patents/figures'

# 专利06-12使用子目录附图
PATENT_FIGURES = {
    '06': [
        ('图1_位置增强财务管理系统架构图.png', '图1 位置增强财务管理系统架构图'),
        ('图2_常驻地点识别流程图.png', '图2 常驻地点识别流程图'),
        ('图3_消费场景识别示意图.png', '图3 消费场景识别示意图'),
        ('图4_地理围栏预警流程图.png', '图4 地理围栏预警流程图'),
        ('图5_位置增强钱龄计算模型图.png', '图5 位置增强钱龄计算模型图'),
    ],
    '07': [
        ('图1_三层去重架构流程图.png', '图1 三层去重架构流程图'),
        ('图2_多因子评分模型示意图.png', '图2 多因子评分模型示意图'),
        ('图3_自适应学习流程图.png', '图3 自适应学习流程图'),
        ('图4_候选集筛选优化示意图.png', '图4 候选集筛选优化示意图'),
    ],
    '08': [
        ('图1_钱龄可视化组件体系示意图.png', '图1 钱龄可视化组件体系示意图'),
        ('图2_多维度下钻交互流程图.png', '图2 多维度下钻交互流程图'),
        ('图3_数据联动机制架构图.png', '图3 数据联动机制架构图'),
        ('图4_智能洞察生成流程图.png', '图4 智能洞察生成流程图'),
    ],
    '09': [
        ('图1_三层信息架构示意图.png', '图1 三层信息架构示意图'),
        ('图2_用户水平评估流程图.png', '图2 用户水平评估流程图'),
        ('图3_渐进加载时序图.png', '图3 渐进加载时序图'),
        ('图4_上下文感知适配示意图.png', '图4 上下文感知适配示意图'),
    ],
    '10': [
        ('图1_智能格式检测流程图.png', '图1 智能格式检测流程图'),
        ('图2_插件化解析器架构图.png', '图2 插件化解析器架构图'),
        ('图3_AI辅助分类流程图.png', '图3 AI辅助分类流程图'),
        ('图4_大文件分片处理流程图.png', '图4 大文件分片处理流程图'),
    ],
    '11': [
        ('图1_离线优先架构示意图.png', '图1 离线优先架构示意图'),
        ('图2_三阶段同步流程图.png', '图2 三阶段同步流程图'),
        ('图3_冲突检测与解决流程图.png', '图3 冲突检测与解决流程图'),
        ('图4_智能同步触发策略图.png', '图4 智能同步触发策略图'),
    ],
    '12': [
        ('图1_隐私保护三层学习架构示意图.png', '图1 隐私保护三层学习架构示意图'),
        ('图2_本地差分隐私处理流程图.png', '图2 本地差分隐私处理流程图'),
        ('图3_安全聚合协议流程图.png', '图3 安全聚合协议流程图'),
        ('图4_模型同步与更新机制图.png', '图4 模型同步与更新机制图'),
    ],
}

PATENT_FILES = {
    '06': '专利06_位置增强财务管理_完整提交版.docx',
    '07': '专利07_多因子交易去重_完整提交版.docx',
    '08': '专利08_财务数据可视化交互_完整提交版.docx',
    '09': '专利09_渐进式披露界面设计_完整提交版.docx',
    '10': '专利10_智能账单解析导入_完整提交版.docx',
    '11': '专利11_离线优先增量同步_完整提交版.docx',
    '12': '专利12_隐私保护协同学习_完整提交版.docx',
}


def add_figures_section(doc, patent_num):
    figures = PATENT_FIGURES.get(patent_num, [])
    if not figures:
        return 0
    
    base_path = os.path.join(FIGURES_DIR, f'patent_{patent_num}')
    
    doc.add_page_break()
    heading = doc.add_heading('说明书附图', level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    figures_added = 0
    for filename, caption in figures:
        filepath = os.path.join(base_path, filename)
        if os.path.exists(filepath):
            doc.add_picture(filepath, width=Inches(5.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            caption_para = doc.add_paragraph()
            caption_run = caption_para.add_run(caption)
            caption_run.bold = True
            caption_run.font.size = Pt(10)
            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph()
            figures_added += 1
        else:
            print(f'    警告: 找不到附图 {filepath}')
    
    return figures_added


def process_patent(patent_num):
    filename = PATENT_FILES.get(patent_num)
    if not filename:
        return
    
    filepath = os.path.join(PATENTS_DIR, filename)
    if not os.path.exists(filepath):
        print(f'  跳过: 文件不存在 {filename}')
        return
    
    print(f'  处理: {filename}')
    
    try:
        doc = Document(filepath)
        figures_added = add_figures_section(doc, patent_num)
        
        if figures_added > 0:
            doc.save(filepath)
            print(f'    成功插入 {figures_added} 张附图')
        else:
            print(f'    未插入任何附图')
    except Exception as e:
        print(f'  错误: {e}')


def main():
    print('开始将附图嵌入专利文档...')
    
    for patent_num in PATENT_FIGURES.keys():
        print(f'专利{patent_num}:')
        process_patent(patent_num)
    
    print('附图嵌入完成!')


if __name__ == '__main__':
    main()
