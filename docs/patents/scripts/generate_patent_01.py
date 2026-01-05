# -*- coding: utf-8 -*-
"""生成专利一：基于FIFO资源池模型的钱龄计算方法及系统"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def create_patent_document():
    doc = Document()

    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # 标题
    title = doc.add_heading('', 0)
    title_run = title.add_run('发明专利申请书')
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # 发明名称
    doc.add_heading('发明名称', level=1)
    doc.add_paragraph('基于FIFO资源池模型的个人财务钱龄计算方法及系统')

    # 技术领域
    doc.add_heading('技术领域', level=1)
    p = doc.add_paragraph()
    p.add_run('[0001] ').bold = True
    p.add_run('本发明涉及个人财务管理技术领域，尤其涉及一种基于FIFO资源池模型的个人财务钱龄计算方法及系统。')

    # 背景技术
    doc.add_heading('背景技术', level=1)

    p = doc.add_paragraph()
    p.add_run('[0002] ').bold = True
    p.add_run('随着移动互联网和移动支付的普及，个人财务管理应用（记账App）已成为用户日常生活的重要工具。然而，现有的记账应用主要关注收支记录和统计分析，存在以下技术问题：')

    p = doc.add_paragraph()
    p.add_run('[0003] ').bold = True
    p.add_run('第一，缺乏资金时间价值���化。现有技术仅能记录和展示消费金额，无法让用户直观理解其消费的时间价值。用户难以知道当前消费的资金是何时获得的收入，无法建立"先赚后花"的健康财务意识。')

    p = doc.add_paragraph()
    p.add_run('[0004] ').bold = True
    p.add_run('第二，财务健康评估维度单一。现有技术的财务健康评估主要基于收支比例、储蓄率等静态指标，缺乏反映资金流动性和财务缓冲能力的动态指标。')

    p = doc.add_paragraph()
    p.add_run('[0005] ').bold = True
    p.add_run('第三，资金来源追溯困难。在多收入源场景下，现有技术无法精确追踪每笔支出具体来自哪笔收入，导致用户财务意识模糊。')

    p = doc.add_paragraph()
    p.add_run('[0006] ').bold = True
    p.add_run('第四，计算效率问题。若采用全量重算方式计算资金时间价值，在大数据量场景下会产生严重的性能问题。')

    p = doc.add_paragraph()
    p.add_run('[0007] ').bold = True
    p.add_run('美国个人理财应用YNAB提出了"Age of Money"（钱龄）概念，但其采用简单的加权平均算法，存在精度不足、无法追溯具体来源等问题。')

    p = doc.add_paragraph()
    p.add_run('[0008] ').bold = True
    p.add_run('因此，需要一种新的技术方案，能够精确计算资金的时间价值，提供可追溯的资金来源链路，并在大数据量场景下保持良好的计算性能。')

    # 发明内容
    doc.add_heading('发明内容', level=1)

    p = doc.add_paragraph()
    p.add_run('[0009] ').bold = True
    p.add_run('本发明要解决的技术问题是：如何精确计算个人财务的"钱龄"指标，实现每笔支出到收入来源的完整追溯，并在���数据量场景下保持高效的增量计算性能。')

    p = doc.add_paragraph()
    p.add_run('[0010] ').bold = True
    p.add_run('为解决上述技术问题，本发明提供一种基于FIFO资源池模型的钱龄计算方法，包括以下步骤：')

    p = doc.add_paragraph()
    p.add_run('[0011] ').bold = True
    p.add_run('步骤S1，资源池创建与管理：')

    p = doc.add_paragraph()
    p.add_run('[0012] ').bold = True
    p.add_run('S1.1 当检测到收入交易时，创建对应的资源池对象，包含以下属性：资源池ID（唯一标识）、收入金额（初始余额）、收入时间戳、当前余额、状态标记（活跃/耗尽）。')

    p = doc.add_paragraph()
    p.add_run('[0013] ').bold = True
    p.add_run('S1.2 将新创建的资源池按时间顺序插入资源池队列。')

    p = doc.add_paragraph()
    p.add_run('[0014] ').bold = True
    p.add_run('S1.3 维护资源池的活跃状态，当余额为零时标记为耗尽。')

    p = doc.add_paragraph()
    p.add_run('[0015] ').bold = True
    p.add_run('步骤S2，FIFO消耗算法：')

    p = doc.add_paragraph()
    p.add_run('[0016] ').bold = True
    p.add_run('S2.1 当检测到支出交易时，获取当前所有活跃状态的资源池。')

    p = doc.add_paragraph()
    p.add_run('[0017] ').bold = True
    p.add_run('S2.2 按资源池创建时间升序排列，遵循先进先出原则。')

    p = doc.add_paragraph()
    p.add_run('[0018] ').bold = True
    p.add_run('S2.3 依次从最早的资源池中扣减金额：对于每个活跃资源池P（按时间升序），如果剩余支出金额小于等于P的当前余额，则从P扣减剩余支出金额，记录消费链路（支出ID、资源池ID、扣减金额、收入时间戳），返回；否则，从P扣减P的当前余额，记录消费链路，将P标记为耗尽，剩余支出金额等于剩余支出金额减去P的当前余额，继续处理下一个资源池。')

    p = doc.add_paragraph()
    p.add_run('[0019] ').bold = True
    p.add_run('S2.4 保存消费链路记录，用于钱龄追溯。')

    p = doc.add_paragraph()
    p.add_run('[0020] ').bold = True
    p.add_run('步骤S3，钱龄计算：')

    p = doc.add_paragraph()
    p.add_run('[0021] ').bold = True
    p.add_run('S3.1 单笔交易钱龄计算：单笔钱龄等于各链路金额乘以链路钱龄之和除以各链路金额之和，其中链路钱龄等于支出时间减去对应资源池收入时间。')

    p = doc.add_paragraph()
    p.add_run('[0022] ').bold = True
    p.add_run('S3.2 账户整体钱龄计算：整体钱龄等于各资源池余额乘以资源池存活天数之和除以各资源池余额之和，其中资源池存活天数等于当前时间减去资源池收入时间。')

    p = doc.add_paragraph()
    p.add_run('[0023] ').bold = True
    p.add_run('步骤S4，增量计算优化：')

    p = doc.add_paragraph()
    p.add_run('[0024] ').bold = True
    p.add_run('S4.1 维护脏数据标记机制：当交易发生变更（新增、修改、删除）时，标记受影响的资源池；仅对标记为脏的资源池进行重算。')

    p = doc.add_paragraph()
    p.add_run('[0025] ').bold = True
    p.add_run('S4.2 增量重算逻辑：确定变更交易的时间点T；获取时间点T之后创建的所有��源池；回滚这些资源池的消费记录；重新执行FIFO消耗算法；更新钱龄计算结果。')

    p = doc.add_paragraph()
    p.add_run('[0026] ').bold = True
    p.add_run('步骤S5，健康等级映射：')

    p = doc.add_paragraph()
    p.add_run('[0027] ').bold = True
    p.add_run('根据计算得到的钱龄值，映射为六级健康等级：60天以上为"非常健康"，30至59天为"健康"，14至29天为"一般"，7至13天为"偏低"，3至6天为"紧张"，0至2天为"月光"。')

    # 附图说明
    doc.add_heading('附图说明', level=1)

    p = doc.add_paragraph()
    p.add_run('[0028] ').bold = True
    p.add_run('图1是本发明实施例提供的FIFO资源池模型架构示意图；')

    p = doc.add_paragraph()
    p.add_run('[0029] ').bold = True
    p.add_run('图2是本发明实施例提供的FIFO消耗算法流程图；')

    p = doc.add_paragraph()
    p.add_run('[0030] ').bold = True
    p.add_run('图3是本发明实施例提供的钱龄计算流程图；')

    p = doc.add_paragraph()
    p.add_run('[0031] ').bold = True
    p.add_run('图4是本发明实施例提供的增量计算优化流程图；')

    p = doc.add_paragraph()
    p.add_run('[0032] ').bold = True
    p.add_run('图5是本发明实施例提供的消费链路追溯示意图。')

    # 具体实施方式
    doc.add_heading('具体实施方式', level=1)

    p = doc.add_paragraph()
    p.add_run('[0033] ').bold = True
    p.add_run('下面结合附图和具体实施例对本发明作进一步详细说明。')

    p = doc.add_paragraph()
    p.add_run('[0034] ').bold = True
    p.add_run('实施例一：基本钱龄计算场景')

    p = doc.add_paragraph()
    p.add_run('[0035] ').bold = True
    p.add_run('假设用户在1月1日收到工资8000元，系统创建资源池Pool#1，属性为：ID=P001，初始金额=8000，收入时间=1月1日，当前余额=8000，状态=活跃。')

    p = doc.add_paragraph()
    p.add_run('[0036] ').bold = True
    p.add_run('用户在1月15日消费2000元，系统执行FIFO消耗：从Pool#1扣减2000元，Pool#1余额变为6000元；记录消费链路：支出ID=E001，资源池ID=P001，金额=2000，资源池收入时间=1月1日。')

    p = doc.add_paragraph()
    p.add_run('[0037] ').bold = True
    p.add_run('计算该笔支出的钱龄：钱龄=（2000×14天）/2000=14天，表示这笔消费花的是14天前赚的钱。')

    p = doc.add_paragraph()
    p.add_run('[0038] ').bold = True
    p.add_run('实施例二：跨资源池消费场景')

    p = doc.add_paragraph()
    p.add_run('[0039] ').bold = True
    p.add_run('假设用户有两个资源池：Pool#1（1月工资，剩余500元，已存活45天）和Pool#2（2月工资，剩余8000元，已存活15天）。')

    p = doc.add_paragraph()
    p.add_run('[0040] ').bold = True
    p.add_run('用户消费1500元，系统执行FIFO消耗：首先从Pool#1扣减500元（耗尽），然后从Pool#2扣减1000元。')

    p = doc.add_paragraph()
    p.add_run('[0041] ').bold = True
    p.add_run('计算该笔支出的钱龄：钱龄=（500×45+1000×15）/1500=25天，采用加权平均计算。')

    p = doc.add_paragraph()
    p.add_run('[0042] ').bold = True
    p.add_run('实施例三：增量计算优化场景')

    p = doc.add_paragraph()
    p.add_run('[0043] ').bold = True
    p.add_run('当用户修改历史交易（如将1月10日的一笔支出金额从500元改为800元）时，系统执行增量重算：标记1月10日之后的所有资源池为"脏"状态；回滚这些资源池的消费记录；从1月10日开始重新执行FIFO消耗算法；更新受影响交易的钱龄值。')

    p = doc.add_paragraph()
    p.add_run('[0044] ').bold = True
    p.add_run('通过增量计算，避免了全量重算，在10000笔交易的场景下，计算时间从2秒优化到50毫秒。')

    # 权利要求书（新页）
    doc.add_page_break()

    title = doc.add_heading('', 0)
    title_run = title.add_run('权利要求书')
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run('1. ').bold = True
    p.add_run('一种基于FIFO资源池模型的个人财务钱龄计算方法，其特征在于，包括以下步骤：')

    p = doc.add_paragraph()
    p.add_run('a) 当检测到收入交易时，创建对应的资源池对象并按时间顺序维护资源池队列；')
    p.paragraph_format.left_indent = Cm(0.5)

    p = doc.add_paragraph()
    p.add_run('b) 当检测到支出交易时，按照先进先出原则依次从最早的资源池中扣减金额，并记录消费链路；')
    p.paragraph_format.left_indent = Cm(0.5)

    p = doc.add_paragraph()
    p.add_run('c) 根据消费链路记��，采用加权平均算法计算单笔交易钱龄和账户整体钱龄；')
    p.paragraph_format.left_indent = Cm(0.5)

    p = doc.add_paragraph()
    p.add_run('d) 当交易发生变更时，通过脏数据标记机制实现增量重算，避免全量计算。')
    p.paragraph_format.left_indent = Cm(0.5)

    p = doc.add_paragraph()
    p.add_run('2. ').bold = True
    p.add_run('根据权利要求1所述的方法，其特征在于，所述资源池对象包含资源池唯一标识、收入金额、收入时间戳、当前余额和状态标记属性。')

    p = doc.add_paragraph()
    p.add_run('3. ').bold = True
    p.add_run('根据权利要求1所述的方法，其特征在于，所述消费链路记录包含支出标识、资源池标识、扣减金额和资源池收入时间戳。')

    p = doc.add_paragraph()
    p.add_run('4. ').bold = True
    p.add_run('根据权利要求1所述的方法，其特征在于，所述单笔交易钱龄的计算公式为：单笔钱龄等于各链路金额乘以链路钱龄之和除以各链路金额之和。')

    p = doc.add_paragraph()
    p.add_run('5. ').bold = True
    p.add_run('根据权利要求1所述的方法，其特征在于，所述账户整体钱龄的计算公式为：整体钱龄等于各资源池余额乘以资源池存活天数之和除以各资源池余额之和。')

    p = doc.add_paragraph()
    p.add_run('6. ').bold = True
    p.add_run('根据权利要求1所述的方法，其特征在于，所述增量重算包括：确定变更交易的时间点，获取该时间点之后的资源池，回滚消费记录，重新执行FIFO消耗算法。')

    p = doc.add_paragraph()
    p.add_run('7. ').bold = True
    p.add_run('根据权利要求1所述的方法，其特征在于，还包括将计算得到的钱龄值映射为健康等级的步骤，所述健康等级分为六级。')

    p = doc.add_paragraph()
    p.add_run('8. ').bold = True
    p.add_run('一种基于FIFO资源池模型的个人财务钱龄计算系统，其特征在于，包括：')

    p = doc.add_paragraph()
    p.add_run('资源池管理模块，用于创建和维���资源池队列；')
    p.paragraph_format.left_indent = Cm(0.5)

    p = doc.add_paragraph()
    p.add_run('FIFO消耗模块，用于按先进先出原则执行资金扣减并记录消费��路；')
    p.paragraph_format.left_indent = Cm(0.5)

    p = doc.add_paragraph()
    p.add_run('钱龄计算模块，用于根据消费链路计算交易钱龄和账户钱龄；')
    p.paragraph_format.left_indent = Cm(0.5)

    p = doc.add_paragraph()
    p.add_run('增量优化模块，用于在交易变更时执行增量重算。')
    p.paragraph_format.left_indent = Cm(0.5)

    # 说明书摘要（新页）
    doc.add_page_break()

    title = doc.add_heading('', 0)
    title_run = title.add_run('说明书摘要')
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run('本发明公开了一种基于FIFO资源池模型的个人财务钱龄计算方法及系统，属于个人财务管理技术领域。该方法包括：当检测到收入交易时创建资源池对象并按时间顺序维护队列；当检测到支出交易时按先进先出原则从资源池扣减金额并记录消费链路；根据消费链路采用加权平均算法计算钱龄；通过脏数据标记机制实现增量重算。本发明解决了现有技术无法量化资金时间价值、无法追溯资金来源、计算效率低等问题，能够精确计算"钱龄"指标，实现每笔支出到收入来源的完整追溯，并在大数据量场景下保持高效的增量计算性能，帮助用户建立健康的财务意识。')

    # 保存文档
    doc.save('D:/code/ai-bookkeeping/docs/patents/专利01_FIFO资源池钱龄计算方法.docx')
    print('专利01文档已生成：专利01_FIFO资源池钱龄计算方法.docx')

if __name__ == '__main__':
    create_patent_document()
