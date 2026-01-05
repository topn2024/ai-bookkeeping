# -*- coding: utf-8 -*-
"""生成专利八：财务数据可视化交互方法及系统"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_paragraph_with_number(doc, number, text):
    p = doc.add_paragraph()
    run = p.add_run(f'[{number:04d}] ')
    run.bold = True
    p.add_run(text)
    return p

def create_patent_document():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    title = doc.add_heading('', 0)
    title_run = title.add_run('发明专利申请书')
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    doc.add_heading('发明名称', level=1)
    doc.add_paragraph('基于钱龄维度的财务数据多层次可视化交互方法及系统')

    doc.add_heading('技术领域', level=1)
    add_paragraph_with_number(doc, 1,
        '本发明涉及数据可视化和人机交互技术领域，尤其涉及一种基于钱龄维度的财务数据多层次可视化交互方法及系统。')

    doc.add_heading('背景技术', level=1)
    add_paragraph_with_number(doc, 2,
        '个人财务管理应用的核心价值在于帮助用户理解和优化自己的财务状况。然而，现有的财务数据可视化技术存在以下问题：')

    add_paragraph_with_number(doc, 3,
        '第一，维度单一。现有技术主要采用收支金额作为唯一可视化维度，通过饼图展示分类占比、柱状图展示时间趋势，无法反映资金的时间价值和流动性状况。')

    add_paragraph_with_number(doc, 4,
        '第二，交互能力有限。大多数财务图表为静态展示，用户无法通过点击、拖拽等交互方式深入探索数据细节，难以发现数据背后的问题和规律。')

    add_paragraph_with_number(doc, 5,
        '第三，图表之间相互独立。不同图表展示不同维度的数据，但彼此之间缺乏联动关系，用户需要在多个页面间跳转才能获得完整的财务视图。')

    add_paragraph_with_number(doc, 6,
        '第四，缺乏智能洞察。现有技术仅展示数据统计结果，无法自动检测异常模式、预测趋势或生成有价值的财务洞察。')

    add_paragraph_with_number(doc, 7,
        '第五，加载体验差。复杂图表的渲染和数据计算导致页面加载缓慢，影响用户体验。')

    add_paragraph_with_number(doc, 8,
        '因此，需要一种新的可视化交互方法，能够以钱龄为核心维度，提供多层次、可交互、联动式的财务数据可视化体验。')

    doc.add_heading('发明内容', level=1)
    add_paragraph_with_number(doc, 9,
        '本发明要解决的技术问题是：如何以"钱龄"为核心创新维度，构建多层次、可交互、具有数据联动能力的财务可视化系统，并集成智能洞察生成能力。')

    add_paragraph_with_number(doc, 10,
        '为解决上述技术问题，本发明提供一种基于钱龄维度的财务数据多层次可视化交互方法，包括以下步骤：')

    add_paragraph_with_number(doc, 11, '步骤S1，钱龄可视化组件体系：')

    add_paragraph_with_number(doc, 12,
        'S1.1 钱龄仪表盘组件：采用渐变色环形图展示当前账户整体钱龄，颜色映射六级健康等级：深绿色（60天以上，非常健康）、绿色（30-59天，健康）、黄色（14-29天，一般）、橙色（7-13天，偏低）、红色（3-6天，紧张）、深红色（0-2天，月光）。仪表盘中心显示具体钱龄天数和较上周的变化趋势。')

    add_paragraph_with_number(doc, 13,
        'S1.2 钱龄趋势图组件：采用面积折线图展示钱龄随时间的变化趋势，X轴为时间（支持30/90/180天切换），Y轴为钱龄天数。图表支持触摸交互，用户点击任意数据点可查看该日钱龄明细及主要影响因素。')

    add_paragraph_with_number(doc, 14,
        'S1.3 资源池瀑布图组件：横向排列展示所有活跃的收入资源池，每个资源池以柱状表示，高度对应剩余金额，颜色对应资源池年龄（越旧颜色越深）。按FIFO顺序从左到右排列，直观展示资金消耗顺序。用户点击单个资源池可查看其消费链路详情。')

    add_paragraph_with_number(doc, 15,
        'S1.4 消费链路桑基图组件：展示特定支出与收入资源池之间的消耗关系，左侧为收入资源池节点，右侧为支出交易节点，中间的流量带表示消耗金额，宽度与金额成正比。')

    add_paragraph_with_number(doc, 16, '步骤S2，多维度下钻交互：')

    add_paragraph_with_number(doc, 17,
        'S2.1 时间维度下钻：支持年→月→周→日四级下钻。用户点击年度统计中的某月，可展开该月详情；点击某日可查看当日所有交易。下钻过程中保持上下文面包屑导航，支持任意层级返回。')

    add_paragraph_with_number(doc, 18,
        'S2.2 分类维度下钻：支持一级分类→二级分类→具体交易三级下钻。在分类饼图中点击"餐饮"扇区，展开显示二级分类（早餐/午餐/晚餐/饮品）；继续点击可查看该子分类的所有交易记录。')

    add_paragraph_with_number(doc, 19,
        'S2.3 钱龄维度下钻：支持整体钱龄→分类钱龄→交易钱龄→资源池追溯四级下钻。用户可从整体钱龄仪表盘下钻到各分类的分项钱龄，再到具体交易的钱龄计算明细，最后追溯该笔支出消耗了哪些收入资源池。')

    add_paragraph_with_number(doc, 20,
        'S2.4 账户维度下钻：支持从多账户汇总视图下钻到单账户详情，查看特定账户的收支情况和钱龄状况。')

    add_paragraph_with_number(doc, 21, '步骤S3，数据联动机制：')

    add_paragraph_with_number(doc, 22,
        'S3.1 选中联动：当用户在图表A中选中某数据元素时，自动高亮图表B、C、D中的关联数据。例如，在分类饼图中选中"餐饮"，趋势图自动切换为餐饮消费趋势，交易列表自动筛选餐饮类交易，资源池图高亮餐饮消费来源。')

    add_paragraph_with_number(doc, 23,
        'S3.2 筛选联动：提供全局筛选器，用户设置时间范围、分类、账户等筛选条件后，页面内所有图表组件同步更新，保持数据一致性。')

    add_paragraph_with_number(doc, 24,
        'S3.3 变更通知机制：当底层数据发生变化（新增交易、编辑交易、删除交易）时，采用发布订阅模式通知所有已订阅的可视化组件进行刷新。为避免频繁刷新影响性能，采用防抖策略，将300毫秒内的连续变更合并为单次刷新。')

    add_paragraph_with_number(doc, 25, '步骤S4，智能洞察生成：')

    add_paragraph_with_number(doc, 26,
        'S4.1 异常检测算法：采用基于历史数据的统计异常检测，包括：钱龄突变检测（日环比下降超过20%）、消费异常检测（单笔金额超过月均值3倍）、频率异常检测（某分类消费频次突增）、时间异常检测（非常规时段的大额消费）。')

    add_paragraph_with_number(doc, 27,
        'S4.2 趋势预测模型：基于历史数据采用时间序列分析预测未来钱龄趋势，当预测结果显示钱龄将在本月内降至警告级别时，提前生成预警洞察。')

    add_paragraph_with_number(doc, 28,
        'S4.3 洞察卡片生成：将检测到的异常和预测结果生成可视化洞察卡片，卡片包含：洞察标题、详细描述、相关数据图表、建议操作。用户可对洞察进行"有帮助"/"无帮助"反馈，用于优化后续洞察生成。')

    add_paragraph_with_number(doc, 29, '步骤S5，性能优化策略：')

    add_paragraph_with_number(doc, 30,
        'S5.1 虚拟滚动：对于长列表（如交易记录列表），采用虚拟滚动技术，仅渲染可视区域内的元素，大幅减少DOM节点数量。')

    add_paragraph_with_number(doc, 31,
        'S5.2 图表懒加载：首屏仅加载核心图表组件，次要图表在滚动至可视区域时才进行加载和渲染。')

    add_paragraph_with_number(doc, 32,
        'S5.3 计算缓存：对于复杂的聚合计算结果（如钱龄计算、分类统计），建立缓存机制，仅在数据变更时增量更新受影响的缓存项。')

    add_paragraph_with_number(doc, 33,
        'S5.4 Web Worker异步计算：将耗时的数据处理任务（如大数据量排序、复杂统计）放入Web Worker中执行，避免阻塞主线程导致界面卡顿。')

    doc.add_heading('附图说明', level=1)
    add_paragraph_with_number(doc, 34, '图1是本发明实施例提供的钱龄可视化组件体系示意图；')
    add_paragraph_with_number(doc, 35, '图2是本发明实施例提供的多维度下钻交互流程图；')
    add_paragraph_with_number(doc, 36, '图3是本发明实施例提供的数据联动机制架构图；')
    add_paragraph_with_number(doc, 37, '图4是本发明实施例提供的智能洞察生成流程图。')

    doc.add_heading('具体实施方式', level=1)
    add_paragraph_with_number(doc, 38, '实施例一：钱龄可视化与下钻交互')

    add_paragraph_with_number(doc, 39,
        '用户打开"分析"页面，系统展示钱龄可视化界面。仪表盘显示当前钱龄42天（健康等级：健康），环形图呈绿色，较上周提升3天（上升箭头）。')

    add_paragraph_with_number(doc, 40,
        '用户点击仪表盘进入钱龄详情页，趋势图展示过去30天钱龄变化曲线。用户点击曲线上1月10日的数据点（当日钱龄35天），系统弹出该日详情卡片，显示当日新增收入（工资8000元）使钱龄上升，以及当日支出（餐饮85元、交通32元）的钱龄影响分析。')

    add_paragraph_with_number(doc, 41,
        '用户继续点击"查看资源池消耗"，进入资源池瀑布图。图中展示4个活跃资源池，1月工资池已耗尽（高度为0），2月工资池剩余500元，3月工资池剩余3000元，4月工资池剩余8000元。')

    add_paragraph_with_number(doc, 42, '实施例二：数据联动交互')

    add_paragraph_with_number(doc, 43,
        '在分析页面，用户点击分类饼图中的"餐饮"扇区（占比35%）。系统触发联动更新：（1）趋势图自动切换为"餐饮消费趋势"，显示餐饮支出的时间变化；（2）右侧交易列表自动筛选显示所有餐饮类交易；（3）资源池图高亮显示被餐饮消费消耗的资源池部分；（4）钱龄仪表盘切换显示"餐饮分类钱龄"。')

    add_paragraph_with_number(doc, 44,
        '用户在全局筛选器中选择"仅显示本周数据"，页面所有图表同步刷新，展示本周的数据视图。')

    add_paragraph_with_number(doc, 45, '实施例三：智能洞察展示')

    add_paragraph_with_number(doc, 46,
        '系统检测到本周餐饮消费较上周增长35%，自动生成洞察卡片："本周餐饮支出异常增长"，详细描述：本周餐饮支出687元，较上周509元增长35%。主要增长来自午餐和外卖。若保持此趋势，本月餐饮预算将超支约20%。建议：考虑增加自带午餐频次，或调整预算分配。')

    add_paragraph_with_number(doc, 47,
        '用户点击"查看详情"可下钻到本周餐饮交易列表，点击"有帮助"为该洞察提供正向反馈。')

    doc.add_heading('有益效果', level=1)
    add_paragraph_with_number(doc, 48, '本发明相比现有技术具有以下有益效果：')
    add_paragraph_with_number(doc, 49, '1. 创新维度：首次将"钱龄"作为核心可视化维度，帮助用户理解资金的时间价值。')
    add_paragraph_with_number(doc, 50, '2. 深度交互：多维度下钻交互使用户可以从概览逐层深入到原始交易，全面理解数据。')
    add_paragraph_with_number(doc, 51, '3. 数据联动：图表间的选中联动和筛选联动提供一致的数据探索体验。')
    add_paragraph_with_number(doc, 52, '4. 智能洞察：自动异常检测和趋势预测帮助用户发现潜在的财务问题。')
    add_paragraph_with_number(doc, 53, '5. 流畅体验：多种性能优化策略确保复杂可视化界面的流畅响应。')

    # 权利要求书
    doc.add_page_break()
    title = doc.add_heading('', 0)
    title_run = title.add_run('权利要求书')
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run('1. ').bold = True
    p.add_run('一种基于钱龄维度的财务数据多层次可视化交互方法，其特征在于，包括以下步骤：')
    for item in [
        'a) 构建以钱龄为核心维度的可视化组件体系，包括钱龄仪表盘、钱龄趋势图、资源池瀑布图和消费链路桑基图；',
        'b) 实现多维度下钻交互，支持时间维度、分类维度、钱龄维度和账户维度的逐层展开探索；',
        'c) 建立图表间的数据联动机制，包括选中联动、筛选联动和变更通知；',
        'd) 集成智能洞察生成模块，自动检测异常模式和预测趋势。'
    ]:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Cm(0.5)

    p = doc.add_paragraph()
    p.add_run('2. ').bold = True
    p.add_run('根据权利要求1所述的方法，其特征在于，所述钱龄仪表盘采用六级健康等级映射，通过渐变色环形图直观展示账户财务健康状态。')

    p = doc.add_paragraph()
    p.add_run('3. ').bold = True
    p.add_run('根据权利要求1所述的方法，其特征在于，所述资源池瀑布图按FIFO顺序横向排列展示活跃资源池，高度表示剩余金额，颜色表示资源池年龄。')

    p = doc.add_paragraph()
    p.add_run('4. ').bold = True
    p.add_run('根据权利要求1所述的方法，其特征在于，所述钱龄维度下钻支持四级：整体钱龄→分类钱龄→交易钱龄→资源池追溯。')

    p = doc.add_paragraph()
    p.add_run('5. ').bold = True
    p.add_run('根据权利要求1所述的方法，其特征在于，所述选中联动在用户选中图表数据元素时，自动高亮其他图表中的关联数据。')

    p = doc.add_paragraph()
    p.add_run('6. ').bold = True
    p.add_run('根据权利要求1所述的方法，其特征在于，所述变更通知机制采用发布订阅模式和防抖策略，将预设时间内的连续变更合并为单次刷新。')

    p = doc.add_paragraph()
    p.add_run('7. ').bold = True
    p.add_run('根据权利要求1所述的方法，其特征在于，所述智能洞察生成包括异常检测和趋势预测，检测类型包括钱龄突变、消费金额异常、频率异常和时间异常。')

    p = doc.add_paragraph()
    p.add_run('8. ').bold = True
    p.add_run('一种基于钱龄维度的财务数据可视化交互系统，其特征在于，包括：')
    for item in [
        '- 组件渲染模块，用于渲染钱龄仪表盘、趋势图、瀑布图等可视化组件；',
        '- 交互处理模块，用于响应用户点击、拖拽等操作，执行下钻导航；',
        '- 数据联动模块，用于管理图表间的选中联动和筛选联动；',
        '- 洞察引擎模块，用于检测异常和生成智能洞察卡片；',
        '- 性能优化模块，用于实现虚拟滚动、懒加载和计算缓存。'
    ]:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Cm(0.5)

    p = doc.add_paragraph()
    p.add_run('9. ').bold = True
    p.add_run('一种计算机可读存储介质，其上存储有计算机程序，其特征在于，所述计算机程序被处理器执行时实现权利要求1至7中任一项所述方法的步骤。')

    # 说明书摘要
    doc.add_page_break()
    title = doc.add_heading('', 0)
    title_run = title.add_run('说明书摘要')
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run('本发明公开了一种基于钱龄维度的财务数据多层次可视化交互方法及系统，属于数据可视化技术领域。该方法以"钱龄"为创新核心维度，构建包括钱龄仪表盘、趋势图、资源池瀑布图和消费链路桑基图的可视化组件体系；实现时间、分类、钱龄、账户四个维度的多层次下钻交互；建立选中联动、筛选联动和变更通知的数据联动机制；集成异常检测和趋势预测的智能洞察生成能力。本发明解决了现有财务可视化维度单一、交互受限、缺乏联动、洞察不足等问题，为用户提供深度、交互式的财务数据探索体验。')

    doc.save('D:/code/ai-bookkeeping/docs/patents/专利08_财务数据可视化交互方法.docx')
    print('专利08文档已生成')

if __name__ == '__main__':
    create_patent_document()
