# -*- coding: utf-8 -*-
"""生成专利十二：隐私保护的协同学习方法"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_paragraph_with_number(doc, number, text):
    p = doc.add_paragraph()
    run = p.add_run(f'[{number:04d}] ')
    run.bold = True
    p.add_run(text)
    return p

def create_patent_document():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    title = doc.add_heading('', 0)
    title_run = title.add_run('发明专利申请书')
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    doc.add_heading('发明名称', level=1)
    doc.add_paragraph('基于差分隐私的分布式协同学习方法及系统')

    doc.add_heading('技术领域', level=1)
    add_paragraph_with_number(doc, 1,
        '本发明涉及机器学习和隐私计算技术领域，尤其涉及一种基于差分隐私的分布式协同学习方法及系统。')

    doc.add_heading('背景技术', level=1)
    add_paragraph_with_number(doc, 2,
        '智能应用需要从用户行为数据中学习以提升准确率和个性化程度。例如，记账应用需要学习用户的分类习惯以提高自动分类准确率。然而，收集和利用用户数据进行机器学习存在以下技术问题：')

    add_paragraph_with_number(doc, 3,
        '第一，隐私泄露风险。传统的集中式机器学习需要将用户原始数据上传到服务器进行训练，存在数据泄露、被滥用或被攻击获取的风险，特别是财务数据属于高度敏感信息。')

    add_paragraph_with_number(doc, 4,
        '第二，模型个性化不足。统一训练的模型只能反映群体平均特征，无法适应个体用户的特殊习惯和偏好，导致对部分用户的预测准确率较低。')

    add_paragraph_with_number(doc, 5,
        '第三，数据孤岛问题。每个用户的数据独立存储在各自设备上，这些分散的数据无法汇聚利用，群体智慧被浪费，新用户无法受益于老用户的学习成果。')

    add_paragraph_with_number(doc, 6,
        '第四，合规挑战。GDPR、《个人信息保护法》等法规对数据收集和使用有严格限制，传统的数据收集方式难以满足合规要求。')

    add_paragraph_with_number(doc, 7,
        '第五，通信效率问题。频繁上传模型更新会消耗大量带宽和电量，不适合移动设备。')

    add_paragraph_with_number(doc, 8,
        '因此，需要一种隐私保护的协同学习方法，能够在保护用户数据隐私的前提下，实现个体学习与群体智慧的融合。')

    doc.add_heading('发明内容', level=1)
    add_paragraph_with_number(doc, 9,
        '本发明要解决的技术问题是：如何在严格保护用户数据隐私的前提下，实现多用户间的协同学习，使个体用户受益于群体智慧，同时保持个性化能力。')

    add_paragraph_with_number(doc, 10,
        '为解决上述技术问题，本发明提供一种基于差分隐私的分布式协同学习方法，包括以下步骤：')

    add_paragraph_with_number(doc, 11, '步骤S1，三层学习架构设计：')

    add_paragraph_with_number(doc, 12,
        'S1.1 个体学习层：每个用户设备本地维护个性化模型，基于用户自己的数据进行训练。模型参数完全存储在本地设备，不外传。个体模型能够学习用户独特的分类习惯、消费模式和语言表达偏好。')

    add_paragraph_with_number(doc, 13,
        'S1.2 协同学习层：周期性聚合多用户的模型更新，提取共性知识。关键原则：不传输原始数据，只传输经过隐私保护处理的模型梯度或规则。聚合后的协同模型包含群体智慧，可帮助个体模型快速优化。')

    add_paragraph_with_number(doc, 14,
        'S1.3 迁移学习层：新用户注册时，继承协同模型作为初始模型，无需从零开始学习。协同模型提供基础能力（如常见商户分类、通用消费模式识别），个体模型在此基础上学习用户特有模式。')

    add_paragraph_with_number(doc, 15, '步骤S2，本地差分隐私保护：')

    add_paragraph_with_number(doc, 16,
        'S2.1 梯度裁剪：在上传模型梯度前，首先进行梯度裁剪，将梯度范数限制在阈值C内。裁剪公式：clipped_gradient = gradient × min(1, C/||gradient||)。这一步限制单个用户对聚合结果的影响上界。')

    add_paragraph_with_number(doc, 17,
        'S2.2 噪声添加：在裁剪后的梯度上添加拉普拉斯噪声或高斯噪声。噪声添加公式：noisy_gradient = clipped_gradient + Noise(0, σ)，其中噪声标准差σ = C × √(2ln(1.25/δ)) / ε，ε为隐私预算参数，δ为松弛参数。')

    add_paragraph_with_number(doc, 18,
        'S2.3 隐私预算管理：每个用户有总隐私预算B（如B=10）。每次参与协同学习消耗预算ε（如ε=0.5）。系统跟踪累计消耗的隐私预算，当累计消耗接近B时，降低参与频率或停止参与协同学习，确保隐私保护强度。')

    add_paragraph_with_number(doc, 19, '步骤S3，数据脱敏规则：')

    add_paragraph_with_number(doc, 20,
        'S3.1 金额脱敏：具体金额转换为金额区间标签（<10元、10-50元、50-200元、200-1000元、>1000元），或转换为相对值（如占月收入比例）。')

    add_paragraph_with_number(doc, 21,
        'S3.2 商户脱敏：具体商户名称泛化为品牌类别或商户类型。例如，"星巴克中关村店"→"咖啡连锁"，"张三理发店"→"个人服务-理发"。')

    add_paragraph_with_number(doc, 22,
        'S3.3 时间脱敏：精确时间模糊化为时段标签（凌晨/早晨/上午/中午/下午/晚间/深夜）和日期类型（工作日/周末/节假日）。')

    add_paragraph_with_number(doc, 23,
        'S3.4 位置脱敏：精确位置泛化为城市级别或区域类型（商业区/居住区/办公区）。')

    add_paragraph_with_number(doc, 24, '步骤S4，安全聚合协议：')

    add_paragraph_with_number(doc, 25,
        'S4.1 聚合流程设计：（1）服务端生成本轮聚合的随机种子，分发给参与用户；（2）各用户使用种子对本地梯度进行加密变换；（3）用户将加密后的梯度上传到服务端；（4）服务端对所有加密梯度进行聚合运算；（5）聚合结果可以正确解密，但单个用户的梯度无法从聚合结果中还原。')

    add_paragraph_with_number(doc, 26,
        'S4.2 最小参与人数要求：每轮聚合需要至少K个用户参与（推荐K≥100），确保单个用户的贡献被充分稀释，无法通过聚合结果推断个体信息。如果参与人数不足K，本轮聚合取消，等待更多用户。')

    add_paragraph_with_number(doc, 27,
        'S4.3 异常检测与剔除：在聚合前检测异常梯度（如范数远超正常值），剔除可能的恶意或错误数据，保护聚合质量。')

    add_paragraph_with_number(doc, 28, '步骤S5，模型同步与增量更新：')

    add_paragraph_with_number(doc, 29,
        'S5.1 同步时机控制：设备端同步策略：WiFi连接且充电状态时执行后台同步；用户主动触发时执行手动同步。服务端发布策略：协同模型每周更新一次；紧急优化可触发即时推送。')

    add_paragraph_with_number(doc, 30,
        'S5.2 增量模型传输：不传输完整模型参数，只传输与上一版本的差异（delta）。差异数据进行压缩编码，进一步减少传输量。客户端收到delta后，在本地合并到现有模型。')

    add_paragraph_with_number(doc, 31,
        'S5.3 模型版本管理：协同模型采用语义化版本号（主版本.次版本.补丁版本）。客户端记录当前使用的模型版本，只下载必要的更新。支持版本回滚，当新版本效果不佳时可恢复。')

    add_paragraph_with_number(doc, 32, '步骤S6，效果评估与反馈闭环：')

    add_paragraph_with_number(doc, 33,
        'S6.1 本地效果评估：在设备端评估模型预测准确率，统计用户修正次数。指标包括：预测准确率、用户修正率、响应时间。评估结果用于判断是否需要更新模型。')

    add_paragraph_with_number(doc, 34,
        'S6.2 匿名效果上报：将效果指标（非原始数据）匿名上报，用于监控协同模型整体表现。上报数据经过聚合处理，无法关联到个人。')

    add_paragraph_with_number(doc, 35,
        'S6.3 持续优化闭环：根据效果反馈调整学习参数（学习率、聚合频率、隐私预算分配），形成持续优化的闭环。')

    doc.add_heading('附图说明', level=1)
    add_paragraph_with_number(doc, 36, '图1是本发明实施例提供的三层学习架构示意图；')
    add_paragraph_with_number(doc, 37, '图2是本发明实施例提供的本地差分隐私处理流程图；')
    add_paragraph_with_number(doc, 38, '图3是本发明实施例提供的安全聚合协议流程图；')
    add_paragraph_with_number(doc, 39, '图4是本发明实施例提供的模型同步与更新机制图。')

    doc.add_heading('具体实施方式', level=1)
    add_paragraph_with_number(doc, 40, '实施例一：商户分类协同学习')

    add_paragraph_with_number(doc, 41,
        '用户A将"瑞幸咖啡"分类为"餐饮/饮品"。本地模型学习到规则：商户名包含"咖啡"或"瑞幸"→分类"餐饮/饮品"。')

    add_paragraph_with_number(doc, 42,
        '系统准备上传学习成果进行协同：原始规则"瑞幸咖啡→餐饮/饮品"经数据脱敏处理为"[咖啡类商户]→餐饮/饮品"。对规则对应的模型梯度进行裁剪和噪声添加，消耗隐私预算ε=0.5。加密后上传到服务端。')

    add_paragraph_with_number(doc, 43,
        '服务端本轮收集到1000个用户的学习成果。安全聚合后发现：90%用户将"咖啡"相关商户归类为"餐饮/饮品"，8%归类为"餐饮/其他"，2%归类为其他。更新协同模型：[咖啡类商户]→餐饮/饮品（置信度90%）。')

    add_paragraph_with_number(doc, 44,
        '新用户B首次使用应用，自动继承协同模型。用户B记录"星巴克"消费，系统根据协同模型自动建议分类"餐饮/饮品"，用户确认。用户B无需经历学习过程，直接受益于群体智慧。')

    add_paragraph_with_number(doc, 45, '实施例二：隐私预算管理')

    add_paragraph_with_number(doc, 46,
        '用户C的隐私预算配置为B=10。使用一年期间，参与了18次协同学习，每次消耗ε=0.5，累计消耗9.0。')

    add_paragraph_with_number(doc, 47,
        '系统检测到剩余预算仅1.0，自动调整策略：降低参与频率（从每周一次改为每月一次）；仅上传高价值学习成果（准确率提升>5%的规则）；提示用户可在设置中重置隐私预算（会清空历史参与记录）。')

    doc.add_heading('有益效果', level=1)
    add_paragraph_with_number(doc, 48, '本发明相比现有技术具有以下有益效果：')
    add_paragraph_with_number(doc, 49, '1. 隐私保护：采用本地差分隐私和安全聚合双重保护，原始数据永不离开用户设备，满足最严格的隐私合规要求。')
    add_paragraph_with_number(doc, 50, '2. 群体智慧：协同学习机制使所有用户受益于群体智慧，新用户无需冷启动即可获得良好预测能力。')
    add_paragraph_with_number(doc, 51, '3. 个性化保持：三层架构中个体学习层独立运行，用户特有习惯得到保留和尊重。')
    add_paragraph_with_number(doc, 52, '4. 可量化隐私：隐私预算机制使隐私保护强度可量化、可控制，用户可自主决定参与程度。')
    add_paragraph_with_number(doc, 53, '5. 通信高效：增量更新和压缩传输大幅减少通信开销，适合移动设备使用。')

    # 权利要求书
    doc.add_page_break()
    title = doc.add_heading('', 0)
    title_run = title.add_run('权利要求书')
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run('1. ').bold = True
    p.add_run('一种基于差分隐私的分布式协同学习方法，其特征在于，包括以下步骤：')
    for item in [
        'a) 构建三层学习架构，包括设备端的个体学习层、服务端的协同学习层和面向新用户的迁移学习层；',
        'b) 在上传模型更新前，执行梯度裁剪和噪声添加的本地差分隐私保护；',
        'c) 对原始数据进行金额、商户、时间、位置的多维度脱敏处理；',
        'd) 采用安全聚合协议聚合多用户的加密梯度，确保单个用户梯度不可还原；',
        'e) 实现隐私预算管理，跟踪和控制用户的隐私暴露程度。'
    ]:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Cm(0.5)

    p = doc.add_paragraph()
    p.add_run('2. ').bold = True
    p.add_run('根据权利要求1所述的方法，其特征在于，所述个体学习层在设备端本地训练和存储个性化模型，模型参数不外传。')

    p = doc.add_paragraph()
    p.add_run('3. ').bold = True
    p.add_run('根据权利要求1所述的方法，其特征在于，所述本地差分隐私保护采用拉普拉斯噪声或高斯噪声，噪声量由隐私预算参数ε和梯度敏感度上界控制。')

    p = doc.add_paragraph()
    p.add_run('4. ').bold = True
    p.add_run('根据权利要求1所述的方法，其特征在于，所述安全聚合要求每轮最小参与用户数不少于预设阈值，确保单个用户贡献被充分稀释。')

    p = doc.add_paragraph()
    p.add_run('5. ').bold = True
    p.add_run('根据权利要求1所述的方法，其特征在于，所述隐私预算管理为每个用户设置总预算上限，每次参与协同学习消耗部分预算，累计接近上限时降低参与频率。')

    p = doc.add_paragraph()
    p.add_run('6. ').bold = True
    p.add_run('根据权利要求1所述的方法，其特征在于，所述迁移学习层使新用户继承协同模型作为初始模型，无需冷启动即可获得基础智能能力。')

    p = doc.add_paragraph()
    p.add_run('7. ').bold = True
    p.add_run('根据权利要求1所述的方法，其特征在于，还包括模型增量更新步骤，仅传输模型差异并进行压缩编码，减少通信开销。')

    p = doc.add_paragraph()
    p.add_run('8. ').bold = True
    p.add_run('一种基于差分隐私的分布式协同学习系统，其特征在于，包括：')
    for item in [
        '- 本地学习模块，用于在设备端训练和维护个体模型；',
        '- 隐私保护模块，用于执行梯度裁剪、噪声添加和数据脱敏；',
        '- 预算管理模块，用于跟踪和控制隐私预算消耗；',
        '- 安全聚合模块，用于在服务端执行加密梯度聚合；',
        '- 模型分发模块，用于将协同模型增量更新分发到客户端；',
        '- 效果评估模块，用于监控模型表现并形成优化闭环。'
    ]:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Cm(0.5)

    p = doc.add_paragraph()
    p.add_run('9. ').bold = True
    p.add_run('一种计算机可读存储介质，其上存储有计算机程序，其特征在于，所述计算机程序被处理器执行时实现权利要求1至7中任一项所述方法的步骤。')

    # 说明书摘要
    doc.add_page_break()
    title = doc.add_heading('', 0)
    title_run = title.add_run('说明书摘要')
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run('本发明公开了一种基于差分隐私的分布式协同学习方法及系统，属于机器学习和隐私计算技术领域。该方法构建个体学习层、协同学习层和迁移学习层的三层架构；在上传模型更新前执行梯度裁剪和噪声添加的本地差分隐私保护；对原始数据进行金额、商户、时间、位置的多维度脱敏；采用安全聚合协议确保单个用户梯度不可还原；实现隐私预算管理控制隐私暴露程度；通过增量更新减少通信开销。本发明解决了传统机器学习中隐私泄露风险、模型个性化不足、数据孤岛、合规挑战等问题，实现了隐私保护与协同学习的平衡。')

    doc.save('D:/code/ai-bookkeeping/docs/patents/专利12_隐私保护协同学习方法.docx')
    print('专利12文档已生成')

if __name__ == '__main__':
    create_patent_document()
