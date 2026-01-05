# -*- coding: utf-8 -*-
"""将所有附图插入到对应的专利文档中"""

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

PATENTS_DIR = 'D:/code/ai-bookkeeping/docs/patents'
FIGURES_DIR = 'D:/code/ai-bookkeeping/docs/patents/figures'

# 专利与附图的映射关系（使用实际文件名）
PATENT_FIGURES = {
    '专利01_FIFO资源池钱龄计算方法.docx': {
        'folder': 'patent_01',
        'figures': [
            '图1_FIFO资源池模型架构示意图.png',
            '图2_FIFO消耗算法流程图.png',
            '图3_钱龄计算流程图.png',
            '图4_增量计算优化流程图.png',
            '图5_消费链路追溯示意图.png',
        ]
    },
    '专利02_多模态融合智能记账识别方法.docx': {
        'folder': 'patent_02',
        'figures': [
            '图1_多模态融合识别系统架构图.png',
            '图2_语音识别处理流程图.png',
            '图3_图像识别处理流程图.png',
            '图4_多笔交易拆分算法流程图.png',
            '图5_实体抽取与分类推荐流程图.png',
        ]
    },
    '专利03_分层自学习与协同学习方法.docx': {
        'folder': 'patent_03',
        'figures': [
            '图1_三层学习架构示意图.png',
            '图2_个体学习流程图.png',
            '图3_协同学习流程图.png',
            '图4_模型融合决策流程图.png',
        ]
    },
    '专利04_零基预算动态分配方法.docx': {
        'folder': 'patent_04',
        'figures': [
            '图1_小金库模型架构图.png',
            '图2_四层零基分配流程图.png',
            '图3_消费拦截流程图.png',
        ]
    },
    '专利05_四维语音交互方法.docx': {
        'folder': 'patent_05',
        'figures': [
            '图1_四维意图分类体系示意图.png',
            '图2_意图识别流程图.png',
            '图3_多轮对话状态机示意图.png',
        ]
    },
    '专利06_位置增强财务管理方法.docx': {
        'folder': 'patent_06',
        'figures': [
            '图1_位置增强财务管理系统架构图.png',
            '图2_常驻地点识别流程图.png',
            '图3_消费场景识别示意图.png',
            '图4_地理围栏预警流程图.png',
            '图5_位置增强钱龄计算模型图.png',
        ]
    },
    '专利07_多因子交易去重方法.docx': {
        'folder': 'patent_07',
        'figures': [
            '图1_三层去重架构流程图.png',
            '图2_多因子评分模型示意图.png',
            '图3_自适应学习流程图.png',
            '图4_候选集筛选优化示意图.png',
        ]
    },
    '专利08_财务数据可视化交互方法.docx': {
        'folder': 'patent_08',
        'figures': [
            '图1_钱龄可视化组件体系示意图.png',
            '图2_多维度下钻交互流程图.png',
            '图3_数据联动机制架构图.png',
            '图4_智能洞察生成流程图.png',
        ]
    },
    '专利09_渐进式披露界面设计方法.docx': {
        'folder': 'patent_09',
        'figures': [
            '图1_三层信息架构示意图.png',
            '图2_用户水平评估流程图.png',
            '图3_渐进加载时序图.png',
            '图4_上下文感知适配示意图.png',
        ]
    },
    '专利10_智能账单解析导入方法.docx': {
        'folder': 'patent_10',
        'figures': [
            '图1_智能格式检测流程图.png',
            '图2_插件化解析器架构图.png',
            '图3_AI辅助分类流程图.png',
            '图4_大文件分片处理流程图.png',
        ]
    },
    '专利11_离线优先增量同步方法.docx': {
        'folder': 'patent_11',
        'figures': [
            '图1_离线优先架构示意图.png',
            '图2_三阶段同步流程图.png',
            '图3_冲突检测与解决流程图.png',
            '图4_智能同步触发策略图.png',
        ]
    },
    '专利12_隐私保护协同学习方法.docx': {
        'folder': 'patent_12',
        'figures': [
            '图1_隐私保护三层学习架构示意图.png',
            '图2_本地差分隐私处理流程图.png',
            '图3_安全聚合协议流程图.png',
            '图4_模型同步与更新机制图.png',
        ]
    },
}


def insert_figures_v3(docx_path, figures_folder, figure_files):
    """在原文档的附图说明后追加图片"""
    doc = Document(docx_path)

    # 找到附图说明和具体实施方式的索引
    figure_section_idx = -1
    impl_section_idx = -1

    for i, para in enumerate(doc.paragraphs):
        if '附图说明' in para.text and 'Heading' in para.style.name:
            figure_section_idx = i
        elif '具体实施方式' in para.text and 'Heading' in para.style.name:
            impl_section_idx = i
            break

    if figure_section_idx == -1:
        print('    [WARN] not found: fu tu shuo ming')
        return False

    inserted_count = 0

    # 获取具体实施方式段落的XML元素
    if impl_section_idx != -1:
        impl_para = doc.paragraphs[impl_section_idx]
        impl_element = impl_para._element

        # 在具体实施方式之前插入图片（逆序插入以保持顺序）
        for figure_file in reversed(figure_files):
            figure_path = os.path.join(FIGURES_DIR, figures_folder, figure_file)
            if not os.path.exists(figure_path):
                print(f'    [SKIP] {figure_file}')
                continue

            # 创建空行段落
            empty_para = doc.add_paragraph()
            impl_element.addprevious(empty_para._element)

            # 创建图片标题段落
            caption_para = doc.add_paragraph()
            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            fig_name = figure_file.replace('.png', '').replace('_', '  ')
            caption_run = caption_para.add_run(fig_name)
            caption_run.font.size = Pt(10)
            impl_element.addprevious(caption_para._element)

            # 创建图片段落
            img_para = doc.add_paragraph()
            img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = img_para.add_run()
            run.add_picture(figure_path, width=Cm(14))
            impl_element.addprevious(img_para._element)

            inserted_count += 1

    doc.save(docx_path)
    return inserted_count


def main():
    print('=' * 60)
    print('Start inserting figures into patent documents...')
    print('=' * 60)

    success_count = 0
    total_figures = 0

    for docx_name, info in PATENT_FIGURES.items():
        docx_path = os.path.join(PATENTS_DIR, docx_name)
        figures_folder = info['folder']
        figure_files = info['figures']

        # Extract patent number for display
        patent_num = docx_name.split('_')[0]
        print(f'\n{patent_num}: {figures_folder} ({len(figure_files)} figures)')

        if not os.path.exists(docx_path):
            print('  [ERROR] Document not found')
            continue

        try:
            count = insert_figures_v3(docx_path, figures_folder, figure_files)
            if count > 0:
                print(f'  [OK] Inserted {count} figures')
                success_count += 1
                total_figures += count
            else:
                print('  [FAIL] No figures inserted')
        except Exception as e:
            print(f'  [ERROR] {e}')

    print('\n' + '=' * 60)
    print(f'Done! {success_count}/{len(PATENT_FIGURES)} documents, {total_figures} figures total')
    print('=' * 60)


if __name__ == '__main__':
    main()
