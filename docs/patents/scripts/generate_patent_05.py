# -*- coding: utf-8 -*-
"""生成专利五：四维语音交互的财务管理方法及系统"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_patent_document():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    title = doc.add_heading('', 0)
    title_run = title.add_run('发明专利申请书')
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    doc.add_heading('发明名称', level=1)
    doc.add_paragraph('四维语音交互的财务管理方法及系统')

    doc.add_heading('技术领域', level=1)
    p = doc.add_paragraph()
    p.add_run('[0001] ').bold = True
    p.add_run('本发明涉及语音交互和个人财务管理技术领域，尤其涉及一种四维语音交互的财务管理方法及系统。')

    doc.add_heading('背景技术', level=1)

    p = doc.add_paragraph()
    p.add_run('[0002] ').bold = True
    p.add_run('语音交互是提升移动应用用户体验的重要方式，但现有财务应用的语音功能存在以下问题：')

    p = doc.add_paragraph()
    p.add_run('[0003] ').bold = True
    p.add_run('第一，功能单一。现有语音功能仅支持简单记账，无法完成查询、设置、导航等复杂操作。')

    p = doc.add_paragraph()
    p.add_run('[0004] ').bold = True
    p.add_run('第二，意图识别能力弱。难以区分用户是要记账、查询还是修改设置，经常误解用户意图。')

    p = doc.add_paragraph()
    p.add_run('[0005] ').bold = True
    p.add_run('第三，缺乏多轮对话能力。当信息不完整时，无法通过追问补全信息，导致交互中断。')

    p = doc.add_paragraph()
    p.add_run('[0006] ').bold = True
    p.add_run('第四，反馈方式单一。仅提供视觉反馈，无法满足驾驶、运动等场景的纯语音交互需求。')

    doc.add_heading('发明内容', level=1)

    p = doc.add_paragraph()
    p.add_run('[0007] ').bold = True
    p.add_run('本发明要解决的技术问题是：如何构建支持记账、配置、导航、查询四维意图的语音交互系统，具备多轮对话能力和语音反馈能力。')

    p = doc.add_paragraph()
    p.add_run('[0008] ').bold = True
    p.add_run('为解决上述技术问题，本发明提供一种四维语音交互的财务管理方法，包括以下步骤：')

    p = doc.add_paragraph()
    p.add_run('[0009] ').bold = True
    p.add_run('步骤S1，四维意图分类体系：')

    p = doc.add_paragraph()
    p.add_run('[0010] ').bold = True
    p.add_run('S1.1 记账意图：用户要记录收支，如"早餐花了15块"、"收到工资8000"。')

    p = doc.add_paragraph()
    p.add_run('[0011] ').bold = True
    p.add_run('S1.2 配置意图：用户要修改设置，如"把默认货币改成美元"、"关闭自动同步"。')

    p = doc.add_paragraph()
    p.add_run('[0012] ').bold = True
    p.add_run('S1.3 导航意图：用户要跳转页面，如"打开预算页面"、"查看本月账单"。')

    p = doc.add_paragraph()
    p.add_run('[0013] ').bold = True
    p.add_run('S1.4 查询意图：用户要查询信息，如"本月花了多少钱"、"我的钱龄是多少"。')

    p = doc.add_paragraph()
    p.add_run('[0014] ').bold = True
    p.add_run('步骤S2，意图识别与分类：')

    p = doc.add_paragraph()
    p.add_run('[0015] ').bold = True
    p.add_run('S2.1 对语音识别结果进行意图分类，采用关键词匹配和语义理解双重机制。')

    p = doc.add_paragraph()
    p.add_run('[0016] ').bold = True
    p.add_run('S2.2 关键词匹配：根据触发词判断意图，如"花了、收到"触发记账意图，"打开、查看"触发导航意图。')

    p = doc.add_paragraph()
    p.add_run('[0017] ').bold = True
    p.add_run('S2.3 语义理解：当关键词匹配不确定时，调用语义模型判断意图。')

    p = doc.add_paragraph()
    p.add_run('[0018] ').bold = True
    p.add_run('步骤S3，多轮对话管理：')

    p = doc.add_paragraph()
    p.add_run('[0019] ').bold = True
    p.add_run('S3.1 维护对话上下文，记录已获取的实体信息和待确认项。')

    p = doc.add_paragraph()
    p.add_run('[0020] ').bold = True
    p.add_run('S3.2 当必要信息缺失时，生成追问话术。如用户说"记一笔午餐"，系统追问"请问金额是多少？"')

    p = doc.add_paragraph()
    p.add_run('[0021] ').bold = True
    p.add_run('S3.3 支持上下文继承，如用户回答"32块"时，自动关联到上一轮的午餐记录。')

    p = doc.add_paragraph()
    p.add_run('[0022] ').bold = True
    p.add_run('步骤S4，语音反馈机制：')

    p = doc.add_paragraph()
    p.add_run('[0023] ').bold = True
    p.add_run('S4.1 执行成功后，生成语音播报确认。如"已记录午餐32元，分类餐饮"。')

    p = doc.add_paragraph()
    p.add_run('[0024] ').bold = True
    p.add_run('S4.2 查询结果通过语音播报返回。如"本月总支出3256元，餐饮占比最高为35%"。')

    p = doc.add_paragraph()
    p.add_run('[0025] ').bold = True
    p.add_run('S4.3 支��语音反馈开关配置，��应不同使用场景。')

    doc.add_heading('附图说明', level=1)

    p = doc.add_paragraph()
    p.add_run('[0026] ').bold = True
    p.add_run('图1是本发明实施例提供的四维意图分类体系示意图；')

    p = doc.add_paragraph()
    p.add_run('[0027] ').bold = True
    p.add_run('图2是本发明实施例提供的意图识别流程图；')

    p = doc.add_paragraph()
    p.add_run('[0028] ').bold = True
    p.add_run('图3是本发明实施例提供的多轮对话状态机示意图。')

    doc.add_heading('具体实施方式', level=1)

    p = doc.add_paragraph()
    p.add_run('[0029] ').bold = True
    p.add_run('实施例一：记账意图多轮对话')

    p = doc.add_paragraph()
    p.add_run('[0030] ').bold = True
    p.add_run('用户："记一笔"。系统识别为记账意图，但缺少金额和描述，追问："请说明消费内容和金额"。用户："午餐"。系统更新上下文（描述=午餐），继续追问："午餐花了多少钱？"用户："32"。系统补全信息，执行记账，语音播报："已记录午餐32元，分类餐饮"。')

    p = doc.add_paragraph()
    p.add_run('[0031] ').bold = True
    p.add_run('实施例二：查询意图处理')

    p = doc.add_paragraph()
    p.add_run('[0032] ').bold = True
    p.add_run('用户："这个月花了多少"。系统识别为查询意图，执行本月支出统计，语音播报："本月总支出4523元，其中餐饮1580元占35%，交通680元占15%，购物1263元占28%"。')

    # 权利要求书
    doc.add_page_break()

    title = doc.add_heading('', 0)
    title_run = title.add_run('权利要求书')
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run('1. ').bold = True
    p.add_run('一种四维语音交互的财务管理方法，其特征在于，包括以下步骤：')

    p = doc.add_paragraph()
    p.add_run('a) 构建四维意图分类体系，包括记账意图、配置意图、导航意图和查询意图；')
    p.paragraph_format.left_indent = Cm(0.5)

    p = doc.add_paragraph()
    p.add_run('b) 对语音输入进行意图识别和分类，采用关键词匹配和语义理解双重机制；')
    p.paragraph_format.left_indent = Cm(0.5)

    p = doc.add_paragraph()
    p.add_run('c) 当必要信息缺失时，通过多轮对话进行追问和信息补全；')
    p.paragraph_format.left_indent = Cm(0.5)

    p = doc.add_paragraph()
    p.add_run('d) 执行完成后通过语音播报进行反馈确认。')
    p.paragraph_format.left_indent = Cm(0.5)

    p = doc.add_paragraph()
    p.add_run('2. ').bold = True
    p.add_run('根据权利要求1所述的方法，其特征在于，所述多轮对话包括上下文维护、追问话术生成和上下文继承机制。')

    p = doc.add_paragraph()
    p.add_run('3. ').bold = True
    p.add_run('一种四维语音交互的财务管理系统，其特征在于，包括：意图识别模块、对话管理模块、意图执行模块和语音反馈模块。')

    # 说明书摘要
    doc.add_page_break()

    title = doc.add_heading('', 0)
    title_run = title.add_run('说明书摘要')
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run('本发明公开了一种四维语音交互的财务管理方法及系统，属于语音交互和个人财务管理技术领域。该方法构建记账、配置、导航、查询四维意图分类体系；采用关键词匹配和语义理解进行意图识别；通过多轮对话管理实现信息追问和补全；执行完成后通过语音播报反馈。本发明解决了现有语音功能单一、意图识别弱、缺乏多轮对话、反馈方式单一等问题，实现了全场景的语音交互财务管理。')

    doc.save('D:/code/ai-bookkeeping/docs/patents/专利05_四维语音交互方法.docx')
    print('专利05文档已生成')

if __name__ == '__main__':
    create_patent_document()
