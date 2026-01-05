# -*- coding: utf-8 -*-
"""生成专利四：零基预算的动态分配与追踪方法及系统"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_patent_document():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    title = doc.add_heading('', 0)
    title_run = title.add_run('发明专利申请书')
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    doc.add_heading('发明名称', level=1)
    doc.add_paragraph('零基预算的动态分配与追踪方法及系统')

    doc.add_heading('技术领域', level=1)
    p = doc.add_paragraph()
    p.add_run('[0001] ').bold = True
    p.add_run('本发明涉及个人财务管理技术领域，尤其涉及一种零基预算的动态分配与追踪方法及系统。')

    doc.add_heading('背景技术', level=1)

    p = doc.add_paragraph()
    p.add_run('[0002] ').bold = True
    p.add_run('预算管理是个人财务管理的核心功能，但现有的预算管理技术存在以下问题：')

    p = doc.add_paragraph()
    p.add_run('[0003] ').bold = True
    p.add_run('第一，固定预算不灵活。传统预算按月固定设置，无法适应收入波动和支出变化。')

    p = doc.add_paragraph()
    p.add_run('[0004] ').bold = True
    p.add_run('第二，分类预算相互孤立。各分类预算独立计算，无法进行灵活调剂。')

    p = doc.add_paragraph()
    p.add_run('[0005] ').bold = True
    p.add_run('第三，超支提醒滞后。仅在超支后提醒，缺乏事前预防机制。')

    p = doc.add_paragraph()
    p.add_run('[0006] ').bold = True
    p.add_run('第四，预算与实际收入脱节。预算金额与实际可支配收入不挂钩，导致预算形同虚设。')

    p = doc.add_paragraph()
    p.add_run('[0007] ').bold = True
    p.add_run('因此，需要一种零基预算方法，能够根据实际收入动态分配预算，支持灵活调剂，并提供消费拦截机制。')

    doc.add_heading('发明内容', level=1)

    p = doc.add_paragraph()
    p.add_run('[0008] ').bold = True
    p.add_run('本发明要解决的技术问题是：如何实现基于实际收入的零基预算，支持分类间灵活调剂，并通过消费拦截机制实现事前预防。')

    p = doc.add_paragraph()
    p.add_run('[0009] ').bold = True
    p.add_run('为解决上述技术问题，本发明提供一种零基预算的动态分配与追踪方法，包括以下步骤：')

    p = doc.add_paragraph()
    p.add_run('[0010] ').bold = True
    p.add_run('步骤S1，小金库模型构建：')

    p = doc.add_paragraph()
    p.add_run('[0011] ').bold = True
    p.add_run('S1.1 将每个预算分类定义为独立的"小金库"，具有独立的余额、预算规则和消费追踪。')

    p = doc.add_paragraph()
    p.add_run('[0012] ').bold = True
    p.add_run('S1.2 每个小金库包含属性：名称、月度预算额度、当前余额、关联的消费分类、是否允许超支。')

    p = doc.add_paragraph()
    p.add_run('[0013] ').bold = True
    p.add_run('步骤S2，零基预算四层分配：')

    p = doc.add_paragraph()
    p.add_run('[0014] ').bold = True
    p.add_run('S2.1 当检测到收入时，启动零基分配流程。')

    p = doc.add_paragraph()
    p.add_run('[0015] ').bold = True
    p.add_run('S2.2 第一层分配固定支出：将收入优先分配给房租、水电、话费等固定周期性支出。')

    p = doc.add_paragraph()
    p.add_run('[0016] ').bold = True
    p.add_run('S2.3 第二层分配必要支出：将剩余收入按比例分配给餐饮、交通等日常必要支出。')

    p = doc.add_paragraph()
    p.add_run('[0017] ').bold = True
    p.add_run('S2.4 第三层分配弹性支出：将剩余收入分配给娱乐、购物等可选支出。')

    p = doc.add_paragraph()
    p.add_run('[0018] ').bold = True
    p.add_run('S2.5 第四层分配储蓄目标：将最终剩余自动分配至储蓄或投资账户。')

    p = doc.add_paragraph()
    p.add_run('[0019] ').bold = True
    p.add_run('步骤S3，消费拦截机制：')

    p = doc.add_paragraph()
    p.add_run('[0020] ').bold = True
    p.add_run('S3.1 当用户记录消费时，检查对应小金库的余额。')

    p = doc.add_paragraph()
    p.add_run('[0021] ').bold = True
    p.add_run('S3.2 如果消费金额超过小金库余额且该小金库设置为不允许超支，触发消费拦截。')

    p = doc.add_paragraph()
    p.add_run('[0022] ').bold = True
    p.add_run('S3.3 展示拦截提示，提供解决方案：从其他小金库调拨资金、修改消费金额、标记为计划外支出。')

    p = doc.add_paragraph()
    p.add_run('[0023] ').bold = True
    p.add_run('步骤S4，跨小金库调拨：')

    p = doc.add_paragraph()
    p.add_run('[0024] ').bold = True
    p.add_run('S4.1 支持用户手动或系统自动在小金库间调拨资金。')

    p = doc.add_paragraph()
    p.add_run('[0025] ').bold = True
    p.add_run('S4.2 自动调拨策略：当某小金库不足时，按优先级从有盈余的小金库自动调拨。')

    p = doc.add_paragraph()
    p.add_run('[0026] ').bold = True
    p.add_run('S4.3 记录调拨历史，支持追溯和分析。')

    doc.add_heading('附图说明', level=1)

    p = doc.add_paragraph()
    p.add_run('[0027] ').bold = True
    p.add_run('图1是本发明实施例提供的小金库模型架构图；')

    p = doc.add_paragraph()
    p.add_run('[0028] ').bold = True
    p.add_run('图2是本发明实施例提供的四层零基分配流程图；')

    p = doc.add_paragraph()
    p.add_run('[0029] ').bold = True
    p.add_run('图3是本发明实施例提供的消费拦截流程图。')

    doc.add_heading('具体实施方式', level=1)

    p = doc.add_paragraph()
    p.add_run('[0030] ').bold = True
    p.add_run('实施例一：零基预算分配场景')

    p = doc.add_paragraph()
    p.add_run('[0031] ').bold = True
    p.add_run('用户月初收到工资10000元，系统执行四层分配：第一层分配固定支出（房租3000元、水电200元、话费100元），小计3300元；第二层分配必要支出（餐饮1500元、交通500元），小计2000元；第三层分配弹性支出（娱乐800元、购物1000元），小计1800元；第四层分配储蓄（剩余2900元自动转入储蓄账户）。')

    p = doc.add_paragraph()
    p.add_run('[0032] ').bold = True
    p.add_run('实施例二：消费拦截场景')

    p = doc.add_paragraph()
    p.add_run('[0033] ').bold = True
    p.add_run('用户餐饮小金库余额80元，尝试记录一笔128元的消费。系统触发拦截提示："餐饮预算不足，本次消费128元，可用余额80元"。提供三个选项：从交通小金库调拨50元（交通余额280元）、修改消费金额为80元、标记为计划外支出。')

    # 权利要求书
    doc.add_page_break()

    title = doc.add_heading('', 0)
    title_run = title.add_run('权利要求书')
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run('1. ').bold = True
    p.add_run('一种零基预算的动态分配与追踪方法，其特征在于，包括以下步骤：')

    p = doc.add_paragraph()
    p.add_run('a) 构建小金库模型，将每个预算分类定义为独立的小金库；')
    p.paragraph_format.left_indent = Cm(0.5)

    p = doc.add_paragraph()
    p.add_run('b) 当检测到收入时，按照固定支出、必要支出、弹性支出、储蓄目标的四层优先级进行零基分配；')
    p.paragraph_format.left_indent = Cm(0.5)

    p = doc.add_paragraph()
    p.add_run('c) 当消费超出小金库余额时，触发消费拦截并提供解决方案；')
    p.paragraph_format.left_indent = Cm(0.5)

    p = doc.add_paragraph()
    p.add_run('d) 支持跨小金库的资金调拨。')
    p.paragraph_format.left_indent = Cm(0.5)

    p = doc.add_paragraph()
    p.add_run('2. ').bold = True
    p.add_run('根据权利要求1所述的方法，其特征在于，所述小金库包含名称、月度预算额度、当前余额、关联分类和超支控制属性。')

    p = doc.add_paragraph()
    p.add_run('3. ').bold = True
    p.add_run('根据权利要求1所述的方法，其特征在于，所述消费拦截的解决方案包括跨小金库调拨、修改金额和标记为计划外支出。')

    p = doc.add_paragraph()
    p.add_run('4. ').bold = True
    p.add_run('一种零基预算的动态分配与追踪系统，其特征在于，包括：小金库管理模块、零基分配模块、消费拦截模块和资金调拨模块。')

    # 说明书摘要
    doc.add_page_break()

    title = doc.add_heading('', 0)
    title_run = title.add_run('说明书摘要')
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run('本发明公开了一种零基预算的动态分配与追踪方法及系统，属于个人财务管理技术领域。该方法构建小金库模型将预算分类为独立账户；收入到账时按固定支出、必要支出、弹性支出、储蓄目标四层优先级进行零基分配；消费时检查小金库余额，不足则触发拦截并提供调拨、修改、标记三种解决方案；支持跨小金库的灵活调剂。本发明解决了传统预算不灵活、分类孤立、超支滞后、与收入脱节等问题，实现了基于实际收入的动态预算管理。')

    doc.save('D:/code/ai-bookkeeping/docs/patents/专利04_零基预算动态分配方法.docx')
    print('专利04文档已生成')

if __name__ == '__main__':
    create_patent_document()
