# -*- coding: utf-8 -*-
"""生成专利六：位置增强的财务管理方法及系统"""

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

def add_paragraph_with_number(doc, number, text, bold_number=True):
    """添加带编号的段落"""
    p = doc.add_paragraph()
    run = p.add_run(f'[{number:04d}] ')
    run.bold = bold_number
    p.add_run(text)
    return p

def create_patent_document():
    doc = Document()

    # 设置页边距
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ==================== 发明专利申请书 ====================
    title = doc.add_heading('', 0)
    title_run = title.add_run('发明专利申请书')
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # 发明名称
    doc.add_heading('发明名称', level=1)
    doc.add_paragraph('位置增强的个人财务管理方法及系统')

    # 技术领域
    doc.add_heading('技术领域', level=1)
    add_paragraph_with_number(doc, 1,
        '本发明涉及地理信息服务和个人财务管理技术领域，尤其涉及一种位置增强的个人财务管理方法及系统。')

    # 背景技术
    doc.add_heading('背景技术', level=1)

    add_paragraph_with_number(doc, 2,
        '随着移动互联网和位置服务的普及，用户的消费行为与地理位置存在密切关联。然而，现有的个人财务管理应用在利用位置信息方面存在以下技术问题：')

    add_paragraph_with_number(doc, 3,
        '第一，消费场景混淆。现有技术无法区分日常消费与出差、旅游期间的异地消费，导致财务分析结果失真。例如，用户出差期间的大额餐饮消费会被计入日常餐饮预算，影响预算执行评估的准确性。')

    add_paragraph_with_number(doc, 4,
        '第二，缺乏主动预警能力。现有技术仅在消费发生后进行记录和统计，无法在用户进入高消费场所前主动提醒预算状况，错失消费决策的最佳干预时机。')

    add_paragraph_with_number(doc, 5,
        '第三，通勤消费识别困难。用户在上下班通勤路线上的消费具有规律性，但现有技术无法自动识别和归类通勤相关消费，影响消费模式分析的精确度。')

    add_paragraph_with_number(doc, 6,
        '第四，钱龄计算不精确。现有的钱龄（资金时间价值）计算方法未考虑消费场景差异，将所有消费同等对待，无法反映用户日常财务状况的真实健康程度。')

    add_paragraph_with_number(doc, 7,
        '第五，位置数据隐私风险。部分应用过度采集位置信息并上传云端，存在用户行踪隐私泄露风险。')

    add_paragraph_with_number(doc, 8,
        '因此，需要一种新的技术方案，能够智能利用位置信息增强财务管理能力，同时保护用户位置隐私。')

    # 发明内容
    doc.add_heading('发明内容', level=1)

    add_paragraph_with_number(doc, 9,
        '本发明要解决的技术问题是：如何利用位置信息增强财务管理能力，实现消费场景智能识别、地理围栏主动预警和位置增强钱龄计算，同时保护用户位置隐私。')

    add_paragraph_with_number(doc, 10,
        '为解决上述技术问题，本发明提供一种位置增强的财务管理方法，包括以下步骤：')

    add_paragraph_with_number(doc, 11,
        '步骤S1，常驻地点智能识别：')

    add_paragraph_with_number(doc, 12,
        'S1.1 位置数据采集：在获取用户授权后，后台周期性采集位置信息，采集频率根据用户活动状态动态调整，静止状态下降低采集频率以节省电量。位置数据采用AES-256加密后仅存储于本地设备。')

    add_paragraph_with_number(doc, 13,
        'S1.2 聚类分析算法：采用DBSCAN密度聚类算法对历史位置进行聚类分析，识别用户的常驻地点。聚类参数包括：邻域半径ε=200米，最小样本数MinPts=10次。')

    add_paragraph_with_number(doc, 14,
        'S1.3 地点标签推断：根据聚类结果的时间特征自动推断地点标签。工作日早9点至晚6点频繁出现的聚类标记为"公司"；夜间和周末频繁出现的聚类标记为"家"；其他高频聚类根据逗留时长和访问频次标记为"常去地点"。')

    add_paragraph_with_number(doc, 15,
        'S1.4 地点学习更新：系统持续学习用户位置模式，当检测到用户搬家或换工作时，自动更新常驻地点标签。更新触发条件为：原常驻地点连续30天未出现，且新位置聚类满足常驻条件。')

    add_paragraph_with_number(doc, 16,
        '步骤S2，消费场景智能识别：')

    add_paragraph_with_number(doc, 17,
        'S2.1 场景分类定义：将消费场景分为三类：（1）日常消费：交易发生位置在常驻地点5公里范围内；（2）异地消费：交易发生位置在常驻地点5公里范围外，且用户当前不处于"旅游模式"；（3）通勤消费：交易发生位置在家和公司之间的通勤路线缓冲区内（缓冲区宽度500米）。')

    add_paragraph_with_number(doc, 18,
        'S2.2 场景自动标记：当用户记录交易时，系统根据交易时间和当前位置自动识别消费场景，并为交易添加场景标签。用户可手动修改场景标签，系统记录修改行为用于优化识别算法。')

    add_paragraph_with_number(doc, 19,
        'S2.3 旅游模式识别：当检测到用户连续24小时处于异地，且移动轨迹符合旅游特征（多个景点POI访问），系统自动提示用户是否开启"旅游模式"。旅游模式下的消费将被标记为"旅游消费"，不计入日常预算统计。')

    add_paragraph_with_number(doc, 20,
        '步骤S3，地理围栏预警系统：')

    add_paragraph_with_number(doc, 21,
        'S3.1 围栏类型定义：系统支持三种类型的地理围栏：（1）高消费商圈围栏：根据历史消费数据自动识别用户在该区域的平均消费较高的商圈；（2）常去消费点围栏：根据用户消费历史自动创建常去餐厅、商店等消费点的围栏；（3）自定义围栏：用户可手动创建任意位置的围栏。')

    add_paragraph_with_number(doc, 22,
        'S3.2 围栏触发机制：采用地理围栏监测技术，当用户进入围栏区域时触发预警检查。检查内容包括：相关预算类别的执行进度、本周/本月在该区域的累计消费、钱龄健康状况。')

    add_paragraph_with_number(doc, 23,
        'S3.3 智能预警策略：根据预算执行状态生成差异化预警：预算使用率<60%时不预警；60%-80%时温和提醒"本月餐饮预算已使用70%"；80%-100%时强提醒"预算即将用尽，建议控制消费"；>100%时警告"预算已超支，本次消费将记为超支"。')

    add_paragraph_with_number(doc, 24,
        'S3.4 预警频率控制：为避免过度打扰，同一围栏24小时内最多触发一次预警；用户可设置免打扰时段；连续忽略3次预警后，该围栏进入静默期。')

    add_paragraph_with_number(doc, 25,
        '步骤S4，位置增强钱龄计算：')

    add_paragraph_with_number(doc, 26,
        'S4.1 分场景钱龄模型：在基础FIFO钱龄算法基础上，引入消费场景维度。系统计算三种钱龄指标：（1）综合钱龄：包含所有消费的整体钱龄；（2）日常钱龄：仅计算日常消费的钱龄，排除异地和旅游消费；（3）分类钱龄：按消费分类计算的细分钱龄。')

    add_paragraph_with_number(doc, 27,
        'S4.2 场景权重调整：在计算综合钱龄时，根据消费场景调整权重。日常消费权重为1.0；通勤消费权重为0.8（因其具有必要性）；异地消费权重为0.5（因其属于偶发性支出）；旅游消费权重为0（不影响日常钱龄评估）。')

    add_paragraph_with_number(doc, 28,
        'S4.3 钱龄健康评估：采用日常钱龄作为用户财务健康的主要评估指标，映射到六级健康等级：危险级（<7天）、警告级（7-14天）、一般级（14-30天）、良好级（30-60天）、优秀级（60-90天）、理想级（>90天）。')

    add_paragraph_with_number(doc, 29,
        '步骤S5，位置数据隐私保护：')

    add_paragraph_with_number(doc, 30,
        'S5.1 本地优先原则：所有位置数据仅存储于用户本地设备，不上传至云端服务器。位置分析算法在设备端执行，确保用户行踪不被第三方获取。')

    add_paragraph_with_number(doc, 31,
        'S5.2 最小化采集：仅在用户主动记账或进入已配置围栏时采集精确位置；后台位置采集采用区域级精度（精确到500米），降低隐私风险。')

    add_paragraph_with_number(doc, 32,
        'S5.3 数据生命周期管理：位置历史数据默认保留90天，超期自动删除。用户可随时查看已存储的位置数据，并可一键清除全部位置历史。')

    add_paragraph_with_number(doc, 33,
        'S5.4 透明控制：在设置页面清晰展示位置权限使用情况，包括：采集频率、存储数据量、各功能的位置使用说明。用户可逐项开关位置相关功能。')

    # 附图说明
    doc.add_heading('附图说明', level=1)

    add_paragraph_with_number(doc, 34, '图1是本发明实施例提供的位置增强财务管理系统架构图；')
    add_paragraph_with_number(doc, 35, '图2是本发明实施例提供的常驻地点识别流程图；')
    add_paragraph_with_number(doc, 36, '图3是本发明实施例提供的消费场景识别示意图；')
    add_paragraph_with_number(doc, 37, '图4是本发明实施例提供的地理围栏预警流程图；')
    add_paragraph_with_number(doc, 38, '图5是本发明实施例提供的位置增强钱龄计算模型图。')

    # 具体实施方式
    doc.add_heading('具体实施方式', level=1)

    add_paragraph_with_number(doc, 39, '下面结合附图和具体实施例对本发明作进一步详细说明。')

    add_paragraph_with_number(doc, 40,
        '实施例一：常驻地点识别与消费场景标记')

    add_paragraph_with_number(doc, 41,
        '用户张先生使用本发明的记账应用一周后，系统通过位置聚类分析识别出两个主要常驻地点：位置A（工作日9:00-18:00高频出现）自动标记为"公司"；位置B（夜间和周末高频出现）自动标记为"家"。')

    add_paragraph_with_number(doc, 42,
        '某工作日中午12:30，张先生在公司附近的商场记录了一笔午餐消费68元。系统检测到消费位置距离"公司"仅800米，自动将该笔消费标记为"日常消费"场景，计入餐饮预算。')

    add_paragraph_with_number(doc, 43,
        '周末，张先生前往另一座城市出差。系统检测到用户位置距离所有常驻地点超过100公里，当张先生记录一笔商务宴请消费580元时，系统自动标记为"异地消费"，并提示用户该笔消费是否计入差旅报销类别。')

    add_paragraph_with_number(doc, 44,
        '实施例二：地理围栏预警')

    add_paragraph_with_number(doc, 45,
        '系统根据张先生的历史消费数据，自动识别出某大型购物中心为"高消费商圈"（历史单次平均消费320元），并创建半径500米的地理围栏。')

    add_paragraph_with_number(doc, 46,
        '某周六下午，张先生前往该购物中心。当其进入围栏区域时，系统触发预警检查，发现本月购物预算已使用85%。系统推送通知："您已进入XX购物中心，本月购物预算剩余15%（¥450），建议理性消费。"')

    add_paragraph_with_number(doc, 47,
        '张先生可选择：查看详细预算情况、暂时关闭该围栏提醒、或忽略继续购物。该提醒帮助张先生在消费决策前获得预算信息，最终其控制消费在300元内，避免超支。')

    add_paragraph_with_number(doc, 48,
        '实施例三：位置增强钱龄计算')

    add_paragraph_with_number(doc, 49,
        '本月张先生的消费构成如下：日常消费8000元、通勤消费600元、出差异地消费2000元、春节旅游消费5000元。采用位置增强钱龄计算：')

    add_paragraph_with_number(doc, 50,
        '综合钱龄计算：各场景消费加权后计算，日常消费（权重1.0）+通勤消费（权重0.8）+异地消费（权重0.5），旅游消费不计入。等效消费金额=8000×1.0+600×0.8+2000×0.5=9480元。')

    add_paragraph_with_number(doc, 51,
        '日常钱龄计算：仅计算日常消费8000元对应的资源池消耗情况，得出日常钱龄为35天，健康等级为"良好"。')

    add_paragraph_with_number(doc, 52,
        '该计算方式避免了旅游大额消费对日常财务健康评估的干扰，更准确地反映了用户的日常理财能力。')

    # 有益效果
    doc.add_heading('有益效果', level=1)

    add_paragraph_with_number(doc, 53, '本发明相比现有技术具有以下有益效果：')

    add_paragraph_with_number(doc, 54,
        '1. 场景感知能力：通过常驻地点识别和消费场景分类，准确区分日常消费、异地消费和通勤消费，提高财务分析的精确度。')

    add_paragraph_with_number(doc, 55,
        '2. 主动预警能力：通过地理围栏技术在用户进入高消费区域前主动提醒预算状况，帮助用户做出更理性的消费决策。')

    add_paragraph_with_number(doc, 56,
        '3. 精准健康评估：位置增强的钱龄计算排除异地和旅游消费的干扰，更准确地反映用户日常财务健康状况。')

    add_paragraph_with_number(doc, 57,
        '4. 隐私安全可控：采用本地存储、最小化采集、透明控制等策略，在利用位置数据的同时充分保护用户隐私。')

    # ==================== 权利要求书 ====================
    doc.add_page_break()

    title = doc.add_heading('', 0)
    title_run = title.add_run('权利要求书')
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # 权利要求1
    p = doc.add_paragraph()
    p.add_run('1. ').bold = True
    p.add_run('一种位置增强的财务管理方法，其特征在于，包括以下步骤：')

    items = [
        'a) 采集用户位置历史数据，采用密度聚类算法识别用户的常驻地点，并根据时间特征自动标记地点类型；',
        'b) 当用户记录交易时，根据交易发生位置与常驻地点的空间关系，识别并标记消费场景，所述消费场景包括日常消费、异地消费和通勤消费；',
        'c) 建立地理围栏监测机制，当用户进入预设围栏区域时，检查相关预算执行状态并生成差异化预警；',
        'd) 在钱龄计算中引入消费场景权重，计算位置增强的综合钱龄和日常钱龄，用于财务健康评估。'
    ]
    for item in items:
        p = doc.add_paragraph()
        p.add_run(item)
        p.paragraph_format.left_indent = Cm(0.5)

    # 权利要求2
    p = doc.add_paragraph()
    p.add_run('2. ').bold = True
    p.add_run('根据权利要求1所述的方法，其特征在于，所述常驻地点识别步骤具体包括：')

    items = [
        '采用DBSCAN密度聚类算法对位置历史进行聚类分析；',
        '根据聚类结果的时间分布特征推断地点类型，工作时间高频出现的聚类标记为公司，非工作时间高频出现的聚类标记为家；',
        '持续监测位置模式变化，当原常驻地点长期未出现时自动更新地点标签。'
    ]
    for item in items:
        p = doc.add_paragraph()
        p.add_run('- ' + item)
        p.paragraph_format.left_indent = Cm(0.5)

    # 权利要求3
    p = doc.add_paragraph()
    p.add_run('3. ').bold = True
    p.add_run('根据权利要求1所述的方法，其特征在于，所述消费场景识别的判定规则为：')

    items = [
        '日常消费：交易位置在常驻地点预设距离范围内；',
        '异地消费：交易位置在常驻地点预设距离范围外；',
        '通勤消费：交易位置在家和公司之间的通勤路线缓冲区内。'
    ]
    for item in items:
        p = doc.add_paragraph()
        p.add_run('- ' + item)
        p.paragraph_format.left_indent = Cm(0.5)

    # 权利要求4
    p = doc.add_paragraph()
    p.add_run('4. ').bold = True
    p.add_run('根据权利要求1所述的方法，其特征在于，所述地理围栏预警机制包括：')

    items = [
        '支持高消费商圈围栏、常去消费点围栏和自定义围栏三种类型；',
        '根据预算使用率采用分级预警策略，包括温和提醒、强提醒和警告三个级别；',
        '设置预警频率控制，同一围栏在预设时间内最多触发一次预警。'
    ]
    for item in items:
        p = doc.add_paragraph()
        p.add_run('- ' + item)
        p.paragraph_format.left_indent = Cm(0.5)

    # 权利要求5
    p = doc.add_paragraph()
    p.add_run('5. ').bold = True
    p.add_run('根据权利要求1所述的方法，其特征在于，所述位置增强钱龄计算包括：')

    items = [
        '为不同消费场景设置权重系数，日常消费权重最高，旅游消费权重为零；',
        '分别计算综合钱龄和日常钱龄，日常钱龄排除异地和旅游消费的影响；',
        '以日常钱龄作为财务健康评估的主要指标。'
    ]
    for item in items:
        p = doc.add_paragraph()
        p.add_run('- ' + item)
        p.paragraph_format.left_indent = Cm(0.5)

    # 权利要求6
    p = doc.add_paragraph()
    p.add_run('6. ').bold = True
    p.add_run('根据权利要求1所述的方法，其特征在于，还包括位置数据隐私保护步骤：')

    items = [
        '位置数据加密后仅存储于用户本地设备，不上传至云端；',
        '采用最小化采集策略，后台采集使用区域级精度；',
        '设置数据保留期限，超期自动删除，用户可随时清除位置历史。'
    ]
    for item in items:
        p = doc.add_paragraph()
        p.add_run('- ' + item)
        p.paragraph_format.left_indent = Cm(0.5)

    # 权利要求7 - 系统权利要求
    p = doc.add_paragraph()
    p.add_run('7. ').bold = True
    p.add_run('一种位置增强的财务管理系统，其特征在于，包括：')

    items = [
        '位置采集模块，用于获取和加密存储用户位置信息；',
        '常驻识别模块，用于通过聚类分析识别和标记用户的常驻地点；',
        '场景识别模块，用于根据交易位置与常驻地点的关系识别消费场景；',
        '围栏预警模块，用于监测地理围栏并在触发时生成预算预警；',
        '钱龄增强模块，用于计算位置增强的综合钱龄和日常钱龄；',
        '隐私保护模块，用于实现数据加密、最小化采集和生命周期管理。'
    ]
    for item in items:
        p = doc.add_paragraph()
        p.add_run('- ' + item)
        p.paragraph_format.left_indent = Cm(0.5)

    # 权利要求8
    p = doc.add_paragraph()
    p.add_run('8. ').bold = True
    p.add_run('根据权利要求7所述的系统，其特征在于，所述场景识别模块还包括旅游模式检测功能，当检测到用户连续处于异地且轨迹符合旅游特征时，自动提示用户开启旅游模式。')

    # 权利要求9
    p = doc.add_paragraph()
    p.add_run('9. ').bold = True
    p.add_run('根据权利要求7所述的系统，其特征在于，所述围栏预警模块支持用户设置免打扰时段，并在用户连续忽略预警后自动进入静默期。')

    # 权利要求10
    p = doc.add_paragraph()
    p.add_run('10. ').bold = True
    p.add_run('一种计算机可读存储介质，其上存储有计算机程序，其特征在于，所述计算机程序被处理器执行时实现权利要求1至6中任一项所述方法的步骤。')

    # ==================== 说明书摘要 ====================
    doc.add_page_break()

    title = doc.add_heading('', 0)
    title_run = title.add_run('说明书摘要')
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run('本发明公开了一种位置增强的个人财务管理方法及系统，属于地理信息服务和个人财务管理技术领域。该方法包括：采集位置历史并通过密度聚类识别常驻地点；根据交易位置与常驻地点的关系识别日常消费、异地消费、通勤消费等场景；建立地理围栏在用户进入高消费区域时主动预警预算状况；在钱龄计算中引入场景权重，计算位置增强的日常钱龄用于财务健康评估。本发明还采用本地存储、最小化采集等策略保护用户位置隐私。本发明解决了现有技术无法区分消费场景、缺乏主动预警、钱龄计算不精确等问题，实现了位置智能增强的个人财务管理。')

    # ==================== 摘要附图 ====================
    doc.add_page_break()

    title = doc.add_heading('', 0)
    title_run = title.add_run('摘要附图')
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # 添加系统架构图（文本形式）
    p = doc.add_paragraph()
    p.add_run('图1：位置增强财务管理系统架构图').bold = True

    code = '''
┌─────────────────────────────────────────────────────────────┐
│                    位置增强财务管理系统                        │
├─────────────────────────────────────────────────────────────┤
│  ┌─────────────┐  ┌─────────────┐  ┌─────────────┐         │
│  │  位置采集    │  │  常驻识别    │  │  场景识别    │         │
│  │    模块     │─→│    模块     │─→│    模块     │         │
│  └─────────────┘  └─────────────┘  └─────────────┘         │
│         │                                   │               │
│         ↓                                   ↓               │
│  ┌─────────────┐                    ┌─────────────┐         │
│  │  围栏预警    │                    │  钱龄增强    │         │
│  │    模块     │                    │    模块     │         │
│  └─────────────┘                    └─────────────┘         │
│         │                                   │               │
│         └───────────────┬───────────────────┘               │
│                         ↓                                   │
│                  ┌─────────────┐                           │
│                  │  隐私保护    │                           │
│                  │    模块     │                           │
│                  └─────────────┘                           │
├─────────────────────────────────────────────────────────────┤
│                      本地加密存储                            │
└─────────────────────────────────────────────────────────────┘
'''

    p = doc.add_paragraph()
    p.add_run(code).font.name = 'Courier New'
    p.add_run(code).font.size = Pt(9)

    # 保存文档
    doc.save('D:/code/ai-bookkeeping/docs/patents/专利06_位置增强财务管理方法.docx')
    print('专利06文档已生成：专利06_位置增强财务管理方法.docx')

if __name__ == '__main__':
    create_patent_document()
