# -*- coding: utf-8 -*-
"""生成专利二：多模态融合的智能记账识别方法及系统"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_patent_document():
    doc = Document()

    # 设置页面边距
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # 标题
    title = doc.add_heading('', 0)
    title_run = title.add_run('发明专利申请书')
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # 发明名称
    doc.add_heading('发明名称', level=1)
    doc.add_paragraph('多模态融合的智能记账识别方法及系统')

    # 技术领域
    doc.add_heading('技术领域', level=1)
    p = doc.add_paragraph()
    p.add_run('[0001] ').bold = True
    p.add_run('本发明涉及人工智能和自然语言处理技术领域，尤其涉及一种多模态融合的智能记账识别方法及系统。')

    # 背景技术
    doc.add_heading('背景技术', level=1)

    p = doc.add_paragraph()
    p.add_run('[0002] ').bold = True
    p.add_run('随着移动支付的普及，用户产生了大量的消费数据，但手动记账仍然是一个繁琐的过程。现有的智能记账技术存在以下问题：')

    p = doc.add_paragraph()
    p.add_run('[0003] ').bold = True
    p.add_run('第一，单一识别通道局限性。现有技术通常只支持单一输入方式（如语音或图片），无法根据用户场景自动选择最优识别方式。')

    p = doc.add_paragraph()
    p.add_run('[0004] ').bold = True
    p.add_run('第二，识别准确率不稳定。在复杂输入场景下（如背景噪音、图片模糊），单一识别引擎的准确率显著下降。')

    p = doc.add_paragraph()
    p.add_run('[0005] ').bold = True
    p.add_run('第三，多笔交易识别困难。当用户一次性输入多笔交易信息时（如"早餐15，午餐32，晚餐48"），现有技术难以准确拆分。')

    p = doc.add_paragraph()
    p.add_run('[0006] ').bold = True
    p.add_run('第四，分类智能化不足。识别出金额后，仍需用户手动选择分类，未能实现端到端的智能记账。')

    p = doc.add_paragraph()
    p.add_run('[0007] ').bold = True
    p.add_run('因此，需要一种多模态融合的识别方法，能够自动选择识别通道、提高准确率、支持多笔交易拆分，并实现自动分类。')

    # 发明内容
    doc.add_heading('发明内容', level=1)

    p = doc.add_paragraph()
    p.add_run('[0008] ').bold = True
    p.add_run('本发明要解决的技术问题是：如何融合语音、图像、文本多种输入模态，实现高准确率的智能记账识别，支持多笔交易拆分和自动分类。')

    p = doc.add_paragraph()
    p.add_run('[0009] ').bold = True
    p.add_run('为解决上述技术问题，本发明提供一种多模态融合的智能记账识别方法，包括以下步骤：')

    p = doc.add_paragraph()
    p.add_run('[0010] ').bold = True
    p.add_run('步骤S1，输入预处理与模态识别：')

    p = doc.add_paragraph()
    p.add_run('[0011] ').bold = True
    p.add_run('S1.1 接收用户输入，自动识别输入模态（语音、图像或文本）。')

    p = doc.add_paragraph()
    p.add_run('[0012] ').bold = True
    p.add_run('S1.2 对语音输入进行降噪预处理和语音活动检测；对图像输入进行旋转校正、对比度增强和区域分割。')

    p = doc.add_paragraph()
    p.add_run('[0013] ').bold = True
    p.add_run('步骤S2，多通道并行识别：')

    p = doc.add_paragraph()
    p.add_run('[0014] ').bold = True
    p.add_run('S2.1 语音通道：调用ASR引擎将语音转换为文本，支持方言识别和连续语音分段。')

    p = doc.add_paragraph()
    p.add_run('[0015] ').bold = True
    p.add_run('S2.2 图像通道：调用OCR引擎识别图像中的文字，根据图像类型（小票、截图、手写）选择专用识别模型。')

    p = doc.add_paragraph()
    p.add_run('[0016] ').bold = True
    p.add_run('S2.3 将各通道识别结果统一转换为文本格式。')

    p = doc.add_paragraph()
    p.add_run('[0017] ').bold = True
    p.add_run('步骤S3，语义理解与实体抽取：')

    p = doc.add_paragraph()
    p.add_run('[0018] ').bold = True
    p.add_run('S3.1 调用大语言模型对文本进行语义理解。')

    p = doc.add_paragraph()
    p.add_run('[0019] ').bold = True
    p.add_run('S3.2 抽取关键实体：金额实体（识别货币符号、数字、大写数字）、时间实体（识别相对时间如"昨天"、绝对时间）、商户实体（识别商户名称、品牌）、分类实体（识别消费类型关键词）。')

    p = doc.add_paragraph()
    p.add_run('[0020] ').bold = True
    p.add_run('步骤S4，多笔交易拆分：')

    p = doc.add_paragraph()
    p.add_run('[0021] ').bold = True
    p.add_run('S4.1 检测输入是否包含多笔交易（通过连词识别、多个金额检测、语义边界分析）。')

    p = doc.add_paragraph()
    p.add_run('[0022] ').bold = True
    p.add_run('S4.2 对多笔交易进行拆分，为每笔交易分配对应的实体信息。')

    p = doc.add_paragraph()
    p.add_run('[0023] ').bold = True
    p.add_run('S4.3 处理实体继承，如当某笔交易缺少时间实体时，继承前一笔交易的时间。')

    p = doc.add_paragraph()
    p.add_run('[0024] ').bold = True
    p.add_run('步骤S5，智能分类与置信度评估：')

    p = doc.add_paragraph()
    p.add_run('[0025] ').bold = True
    p.add_run('S5.1 基于用户历史分类习惯、商户库匹配、语义推断三个维度，为每笔交易推荐分类。')

    p = doc.add_paragraph()
    p.add_run('[0026] ').bold = True
    p.add_run('S5.2 计算每个识别结果的置信度，当置信度低于阈值时标记为需人工确认。')

    p = doc.add_paragraph()
    p.add_run('[0027] ').bold = True
    p.add_run('步骤S6，结果融合与输出：')

    p = doc.add_paragraph()
    p.add_run('[0028] ').bold = True
    p.add_run('S6.1 将识别结果结构化为标准交易对象（金额、日期、分类、备注、置信度）。')

    p = doc.add_paragraph()
    p.add_run('[0029] ').bold = True
    p.add_run('S6.2 对高置信度结果自动入账，对低置信度结果提示用户确认。')

    # 附图说明
    doc.add_heading('附图说明', level=1)

    p = doc.add_paragraph()
    p.add_run('[0030] ').bold = True
    p.add_run('图1是本发明实施例提供的多模态融合识别系统架构图；')

    p = doc.add_paragraph()
    p.add_run('[0031] ').bold = True
    p.add_run('图2是本发明实施例提供的语音识别处理流程图；')

    p = doc.add_paragraph()
    p.add_run('[0032] ').bold = True
    p.add_run('图3是本发明实施例提供的图像识别处理流程图；')

    p = doc.add_paragraph()
    p.add_run('[0033] ').bold = True
    p.add_run('图4是本发明实施例提供的多笔交易拆分算法流程图；')

    p = doc.add_paragraph()
    p.add_run('[0034] ').bold = True
    p.add_run('图5是本发明实施例提供的实体抽取与分类推荐流程图。')

    # 具体实施方式
    doc.add_heading('具体实施方式', level=1)

    p = doc.add_paragraph()
    p.add_run('[0035] ').bold = True
    p.add_run('下面结合附图和具体实施例对本发明作进一步详细说明。')

    p = doc.add_paragraph()
    p.add_run('[0036] ').bold = True
    p.add_run('实施例一：语音输入场景')

    p = doc.add_paragraph()
    p.add_run('[0037] ').bold = True
    p.add_run('用户说"早餐15，打车到公司花了32"，系统处理流程如下：')

    p = doc.add_paragraph()
    p.add_run('[0038] ').bold = True
    p.add_run('语音预处理：进行降噪、端点检测；ASR识别：将语音转换为文本"早餐15，打车到公司花了32"；多笔交易检测：检测到顿号分隔符和两个金额，判定为多笔交易；交易拆分：拆分为交易1（早餐、15元）和交易2（打车、32元）；智能分类：根据"早餐"推荐分类"餐饮"，根据"打车"推荐分类"交通"；输出结果：生成两笔待确认交易。')

    p = doc.add_paragraph()
    p.add_run('[0039] ').bold = True
    p.add_run('实施例二：图像输入场景')

    p = doc.add_paragraph()
    p.add_run('[0040] ').bold = True
    p.add_run('用户拍摄超市小票，系统处理流程如下：')

    p = doc.add_paragraph()
    p.add_run('[0041] ').bold = True
    p.add_run('图像预处理：旋转校正、对比度增强；图像类型识别：检测为"小票"类型，调用小票专用OCR模型；文字识别：识别商品列表和总金额；结构化抽取：提取商户名称"家乐福"、商品明细、总金额"158.50"、交易时间；分类推荐：根据"家乐福"匹配商户库，推荐分类"购物/超市"；输出结果：生成一笔包含商品明细的交易记录。')

    p = doc.add_paragraph()
    p.add_run('[0042] ').bold = True
    p.add_run('实施例三：低置信度处理场景')

    p = doc.add_paragraph()
    p.add_run('[0043] ').bold = True
    p.add_run('用户上传一张模糊的支付截图，系统识别到金额"?8.00"（首位数字不确定），置信度仅为65%。系统展示识别结果并标记金额需要确认，用户确认为"38.00"后完成入账。')

    # 权利要求书
    doc.add_page_break()

    title = doc.add_heading('', 0)
    title_run = title.add_run('权利要求书')
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run('1. ').bold = True
    p.add_run('一种多模态融合的智能记账识别方法，其特征在于，包括以下步骤：')

    p = doc.add_paragraph()
    p.add_run('a) 接收用户输入并自动识别输入模态，对输入进行预处理；')
    p.paragraph_format.left_indent = Cm(0.5)

    p = doc.add_paragraph()
    p.add_run('b) 根据输入模态调用对应的识别引擎，将语音或图像转换为文本；')
    p.paragraph_format.left_indent = Cm(0.5)

    p = doc.add_paragraph()
    p.add_run('c) 调用语义理解模型对文本进行实体抽取，包括金额、时间、商户、分类实体；')
    p.paragraph_format.left_indent = Cm(0.5)

    p = doc.add_paragraph()
    p.add_run('d) 检测并拆分多笔交易，为每笔交易分配实体信息；')
    p.paragraph_format.left_indent = Cm(0.5)

    p = doc.add_paragraph()
    p.add_run('e) 基于用户历史、商户库、语义推断进行智能分类并评估置信度；')
    p.paragraph_format.left_indent = Cm(0.5)

    p = doc.add_paragraph()
    p.add_run('f) 输出结构化交易记录，高置信度自动入账，低置信度提示确认。')
    p.paragraph_format.left_indent = Cm(0.5)

    p = doc.add_paragraph()
    p.add_run('2. ').bold = True
    p.add_run('根据权利要求1所述的方法，其特征在于，所述输入预处理包括：对语音进行降噪和语音活动检测，对图像进行旋转校正和对比度增强。')

    p = doc.add_paragraph()
    p.add_run('3. ').bold = True
    p.add_run('根据权利要求1所述的方法，其特征在于，所述图像识别根据图像类型自动选择专用模型，所述图像类型包括小票、支付截图和手写笔记。')

    p = doc.add_paragraph()
    p.add_run('4. ').bold = True
    p.add_run('根据权利要求1所述的方法，其特征在于，所述多笔交易检测通过连词识别、多金额检测和语义边界分析实现。')

    p = doc.add_paragraph()
    p.add_run('5. ').bold = True
    p.add_run('根据权利要求1所述的方法，其特征在于，所述交易拆分还包括实体继承处理，当某笔交易缺少特定实体时继承相邻交易的对应实体。')

    p = doc.add_paragraph()
    p.add_run('6. ').bold = True
    p.add_run('根据权利要求1所述的方法，其特征在于，所述智能分类基于用户历史分类习惯、商户库匹配和语义推断三个维度的加权评分。')

    p = doc.add_paragraph()
    p.add_run('7. ').bold = True
    p.add_run('一种多模态融合的智能记账识别系统，其特征在于，包括：输入预处理模块、多通道识别模块、语义理解模块、交易拆分模块、智能分类模块和结果输出模块。')

    # 说明书摘要
    doc.add_page_break()

    title = doc.add_heading('', 0)
    title_run = title.add_run('说明书摘要')
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run('本发明公开了一种多模态融合的智能记账识别方法及系统，属于人工智能和自然语言处理技术领域。该方法包括：自动识别输入模态并预处理；调用对应识别引擎将语音或图像转换为文本；通过语义理解模型抽取金额、时间、商户、分类等实体；检测并拆分多笔交易；基于用户历史、商户库、语义推断进行智能分类；根据置信度决定自动入账或提示确认。本发明解决了现有技术单一识别通道局限、准确率不稳定、多笔交易难拆分、分类不智能等问题，实现了端到端的智能记账体验。')

    # 保存
    doc.save('D:/code/ai-bookkeeping/docs/patents/专利02_多模态融合智能记账识别方法.docx')
    print('专利02文档已生成')

if __name__ == '__main__':
    create_patent_document()
