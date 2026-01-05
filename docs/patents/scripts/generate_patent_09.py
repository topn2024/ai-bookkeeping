# -*- coding: utf-8 -*-
"""生成专利九：渐进式披露的移动应用界面设计方法"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_paragraph_with_number(doc, number, text):
    p = doc.add_paragraph()
    run = p.add_run(f'[{number:04d}] ')
    run.bold = True
    p.add_run(text)
    return p

def create_patent_document():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    title = doc.add_heading('', 0)
    title_run = title.add_run('发明专利申请书')
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    doc.add_heading('发明名称', level=1)
    doc.add_paragraph('基于渐进式披露原则的移动应用界面自适应设计方法及系统')

    doc.add_heading('技术领域', level=1)
    add_paragraph_with_number(doc, 1,
        '本发明涉及人机交互和用户界面设计技术领域，尤其涉及一种基于渐进式披露原则的移动应用界面自适应设计方法及系统。')

    doc.add_heading('背景技术', level=1)
    add_paragraph_with_number(doc, 2,
        '随着移动应用功能日益丰富，如何在有限的屏幕空间内合理组织和展示信息成为用户体验设计的核心挑战。现有的移动应用界面设计存在以下问题：')

    add_paragraph_with_number(doc, 3,
        '第一，信息过载问题。功能丰富的应用往往一次性展示大量功能入口和数据信息，导致用户认知负担过重，难以快速找到所需功能。')

    add_paragraph_with_number(doc, 4,
        '第二，新手不友好。复杂的界面布局和专业术语让新用户感到困惑，增加学习成本，导致用户流失率上升。')

    add_paragraph_with_number(doc, 5,
        '第三，高级用户效率受限。为照顾新手用户而过度简化的界面，反而影响了高级用户的操作效率，无法满足其快速访问高级功能的需求。')

    add_paragraph_with_number(doc, 6,
        '第四，加载体验差。页面内容一次性加载导致首屏渲染时间过长，用户在等待过程中看到空白页面或加载动画，体验不佳。')

    add_paragraph_with_number(doc, 7,
        '第五，缺乏场景适配。界面内容固定不变，无法根据用户当前场景（如时间、位置、使用习惯）动态调整展示内容。')

    add_paragraph_with_number(doc, 8,
        '因此，需要一种渐进式披露的界面设计方法，能够根据用户水平和使用场景，自适应地展示合适的信息层次。')

    doc.add_heading('发明内容', level=1)
    add_paragraph_with_number(doc, 9,
        '本发明要解决的技术问题是：如何设计一种能够根据用户水平和使用场景，自适应展示信息层次的移动应用界面系统，兼顾新手友好和高级用户效率。')

    add_paragraph_with_number(doc, 10,
        '为解决上述技术问题，本发明提供一种基于渐进式披露原则的移动应用界面自适应设计方法，包括以下步骤：')

    add_paragraph_with_number(doc, 11, '步骤S1，三层信息架构设计：')

    add_paragraph_with_number(doc, 12,
        'S1.1 核心信息层（第一层）：包含用户必须看到且高频使用的信息和功能。以记账应用为例，包括：账户余额概览、当日收支快捷入口、快捷记账按钮。该层信息对所有用户100%可见，无条件展示。')

    add_paragraph_with_number(doc, 13,
        'S1.2 常用信息层（第二层）：包含用户经常需要但非每次必看的信息。包括：本周/本月收支统计、分类消费概览、近期交易列表（最新5条）、预算执行进度。该层信息根据用户水平决定展示方式。')

    add_paragraph_with_number(doc, 14,
        'S1.3 详细信息层（第三层）：包含偶尔需要的高级功能和详细数据。包括：完整交易历史、详细分析报告、高级设置选项、数据导入导出、多账户管理。该层信息默认收起，按需展开。')

    add_paragraph_with_number(doc, 15, '步骤S2，用户水平动态评估：')

    add_paragraph_with_number(doc, 16,
        'S2.1 行为特征采集：采集用户使用行为数据，包括：使用天数、功能访问频次、操作复杂度（使用过哪些高级功能）、停留时长分布、操作错误率。')

    add_paragraph_with_number(doc, 17,
        'S2.2 水平等级划分：根据行为特征将用户划分为三个等级。新手用户（使用时长<7天或操作复杂度<30%）：优先展示第一层，引导式披露第二层，隐藏第三层。普通用户（7-30天且操作复杂度30%-70%）：展示第一、二层，按需展示第三层。高级用户（>30天且操作复杂度>70%）：直接展示全部层次，提供快捷键和高级功能入口。')

    add_paragraph_with_number(doc, 18,
        'S2.3 动态升级机制：用户水平不是静态标签，系统持续监测用户行为，当用户满足升级条件时自动调整界面展示层次。同时支持手动调整，用户可在设置中选择"简洁模式"或"专业模式"。')

    add_paragraph_with_number(doc, 19, '步骤S3，骨架屏与渐进加载：')

    add_paragraph_with_number(doc, 20,
        'S3.1 智能骨架屏生成：根据目标页面布局自动生成对应的骨架屏。骨架屏采用灰色渐变动画，模拟真实布局的占位效果。支持列表、卡片、图表、头像等多种组件类型的骨架模板。')

    add_paragraph_with_number(doc, 21,
        'S3.2 四级优先级加载策略：优先级1（关键路径）：视口内核心内容，目标加载时间<500ms，包括账户余额、当日收支等；优先级2（增强体验）：视口内次要内容，目标1秒内加载，包括统计图表、消息提示等；优先级3（预加载）：视口外但可能需要的内容，在用户滚动前预加载；优先级4（懒加载）：按需加载内容，仅在用户触发时加载。')

    add_paragraph_with_number(doc, 22,
        'S3.3 加载状态管理：为每个组件定义三种状态：loading（显示骨架屏）、error（显示错误提示和重试按钮）、success（显示真实内容）。状态转换动画采用淡入效果，避免突兀的内容闪烁。')

    add_paragraph_with_number(doc, 23, '步骤S4，上下文感知的内容适配：')

    add_paragraph_with_number(doc, 24,
        'S4.1 时间感知：根据当前时间动态调整展示内容。早间（6-9点）：突出早餐记账入口和今日预算概览；午间（11-14点）：突出午餐消费记录和工作日消费统计；月末：突出预算执行总结和下月预算规划入口；发薪日：突出收入确认和预算分配向导。')

    add_paragraph_with_number(doc, 25,
        'S4.2 使用习惯感知：学习用户操作习惯，将高频功能提升到更显眼的位置。例如，用户经常在晚间记录当日消费，则晚间自动展开"今日消费汇总"卡片。')

    add_paragraph_with_number(doc, 26,
        'S4.3 数据状态感知：根据用户数据状态调整展示。新用户无数据时展示引导内容；预算超支时突出预警卡片；钱龄健康时展示鼓励内容。')

    add_paragraph_with_number(doc, 27, '步骤S5，折叠展开交互设计：')

    add_paragraph_with_number(doc, 28,
        'S5.1 智能折叠规则：第二层和第三层内容默认折叠，显示摘要信息和"展开"提示。折叠状态下显示关键指标（如"本月支出¥3,256"），展开后显示详细图表和列表。')

    add_paragraph_with_number(doc, 29,
        'S5.2 展开偏好记忆：记录用户对各组件的展开偏好。若用户连续3次主动展开某组件，则该组件对该用户默认展开。若用户从不展开某组件，考虑在该用户界面中降低其显示优先级。')

    add_paragraph_with_number(doc, 30,
        'S5.3 动画过渡效果：折叠/展开过程采用平滑动画（时长200-300ms），内容高度变化采用弹性缓动曲线，避免界面突兀跳动。')

    doc.add_heading('附图说明', level=1)
    add_paragraph_with_number(doc, 31, '图1是本发明实施例提供的三层信息架构示意图；')
    add_paragraph_with_number(doc, 32, '图2是本发明实施例提供的用户水平评估流程图；')
    add_paragraph_with_number(doc, 33, '图3是本发明实施例提供的渐进加载时序图；')
    add_paragraph_with_number(doc, 34, '图4是本发明实施例提供的上下文感知适配示意图。')

    doc.add_heading('具体实施方式', level=1)
    add_paragraph_with_number(doc, 35, '实施例一：新用户首次使用体验')

    add_paragraph_with_number(doc, 36,
        '新用户张先生首次打开记账应用，系统根据新用户状态执行以下渐进式披露策略：')

    add_paragraph_with_number(doc, 37,
        '首屏加载：200ms内显示骨架屏占位，500ms内完成核心信息加载，首屏显示简洁的欢迎界面和"记一笔"大按钮。第二层的统计信息显示为折叠状态，仅显示"暂无数据，记录第一笔开始吧"的引导文案。第三层的高级功能入口完全隐藏，底部导航仅显示"首页""记账""我的"三个入口。')

    add_paragraph_with_number(doc, 38,
        '张先生记录了第一笔消费后，系统渐进披露第二层信息：展开显示"今日支出"卡片，并用气泡提示引导用户了解统计功能。')

    add_paragraph_with_number(doc, 39, '实施例二：用户水平动态升级')

    add_paragraph_with_number(doc, 40,
        '用户李女士使用应用满7天，累计记账35笔，使用过分类管理功能。系统判定其从"新手"升级为"普通用户"：')

    add_paragraph_with_number(doc, 41,
        '界面自动调整：首页默认展开"本周统计"卡片；底部导航增加"分析"入口；设置页面显示更多配置选项。同时推送引导提示："您已解锁更多功能，点击查看"。')

    add_paragraph_with_number(doc, 42, '实施例三：上下文感知适配')

    add_paragraph_with_number(doc, 43,
        '某月最后一周，系统检测到月末时间节点，自动调整界面展示：首页顶部插入"月度预算执行进度"卡片，显示各分类预算使用百分比；为预算超支的分类显示红色预警标识；在"我的"页面突出"生成月度报告"功能入口。')

    doc.add_heading('有益效果', level=1)
    add_paragraph_with_number(doc, 44, '本发明相比现有技术具有以下有益效果：')
    add_paragraph_with_number(doc, 45, '1. 降低认知负担：三层信息架构和渐进披露策略避免信息过载，新用户上手难度降低60%。')
    add_paragraph_with_number(doc, 46, '2. 兼顾不同用户：自适应机制使新手和高级用户都能获得合适的界面体验。')
    add_paragraph_with_number(doc, 47, '3. 提升加载体验：骨架屏和渐进加载使首屏感知加载时间缩短至500ms以内。')
    add_paragraph_with_number(doc, 48, '4. 场景适配：上下文感知使界面内容更贴合用户当前需求。')
    add_paragraph_with_number(doc, 49, '5. 个性化记忆：展开偏好记忆使界面逐渐适应个体使用习惯。')

    # 权利要求书
    doc.add_page_break()
    title = doc.add_heading('', 0)
    title_run = title.add_run('权利要求书')
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run('1. ').bold = True
    p.add_run('一种基于渐进式披露原则的移动应用界面自适应设计方法，其特征在于，包括以下步骤：')
    for item in [
        'a) 将界面信息划分为核心信息层、常用信息层和详细信息层三个层次，各层具有不同的展示优先级；',
        'b) 采集用户使用行为数据，根据使用时长和操作复杂度评估用户水平，划分为新手、普通和高级三个等级；',
        'c) 根据用户水平等级自适应选择信息展示层次和方式；',
        'd) 采用骨架屏和多级优先加载策略，优化页面加载体验；',
        'e) 根据时间、使用习惯和数据状态进行上下文感知的内容适配。'
    ]:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Cm(0.5)

    p = doc.add_paragraph()
    p.add_run('2. ').bold = True
    p.add_run('根据权利要求1所述的方法，其特征在于，所述核心信息层包含必看且高频使用的信息，对所有用户100%可见；常用信息层根据用户水平决定展示方式；详细信息层默认收起，按需展开。')

    p = doc.add_paragraph()
    p.add_run('3. ').bold = True
    p.add_run('根据权利要求1所述的方法，其特征在于，所述用户水平评估采用使用天数和操作复杂度两个维度的综合评估，并支持动态升级和手动调整。')

    p = doc.add_paragraph()
    p.add_run('4. ').bold = True
    p.add_run('根据权利要求1所述的方法，其特征在于，所述多级优先加载包括四个优先级：关键路径内容优先加载、增强体验内容次优先、视口外内容预加载、按需内容懒加载。')

    p = doc.add_paragraph()
    p.add_run('5. ').bold = True
    p.add_run('根据权利要求1所述的方法，其特征在于，所述骨架屏根据目标页面布局动态生成，支持列表、卡片、图表等多种组件类型。')

    p = doc.add_paragraph()
    p.add_run('6. ').bold = True
    p.add_run('根据权利要求1所述的方法，其特征在于，还包括折叠展开交互设计步骤，记录用户的展开偏好，对用户连续多次主动展开的组件自动默认展开。')

    p = doc.add_paragraph()
    p.add_run('7. ').bold = True
    p.add_run('一种基于渐进式披露原则的移动应用界面系统，其特征在于，包括：')
    for item in [
        '- 信息分层模块，用于管理三层信息架构和各层内容配置；',
        '- 用户评估模块，用于采集行为数据并评估用户水平等级；',
        '- 自适应展示模块，用于根据用户水平决定信息展示策略；',
        '- 渐进加载模块，用于实现骨架屏和多级优先加载；',
        '- 上下文感知模块，用于根据场景动态调整展示内容。'
    ]:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Cm(0.5)

    p = doc.add_paragraph()
    p.add_run('8. ').bold = True
    p.add_run('一种计算机可读存储介质，其上存储有计算机程序，其特征在于，所述计算机程序被处理器执行时实现权利要求1至6中任一项所述方法的步骤。')

    # 说明书摘要
    doc.add_page_break()
    title = doc.add_heading('', 0)
    title_run = title.add_run('说明书摘要')
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run('本发明公开了一种基于渐进式披露原则的移动应用界面自适应设计方法及系统，属于人机交互技术领域。该方法将界面信息划分为核心、常用、详细三个层次；采集用户行为数据评估用户水平，划分新手、普通、高级三个等级；根据用户水平自适应选择信息展示层次；采用骨架屏和四级优先加载策略优化加载体验；根据时间、使用习惯和数据状态进行上下文感知的内容适配。本发明解决了现有应用信息过载、新手不友好、高级用户效率受限、加载体验差等问题，实现了兼顾不同用户群体的自适应界面设计。')

    doc.save('D:/code/ai-bookkeeping/docs/patents/专利09_渐进式披露界面设计方法.docx')
    print('专利09文档已生成')

if __name__ == '__main__':
    create_patent_document()
