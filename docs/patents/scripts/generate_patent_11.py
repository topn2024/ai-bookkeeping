# -*- coding: utf-8 -*-
"""生成专利十一：移动端离线优先的增量同步方法"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_paragraph_with_number(doc, number, text):
    p = doc.add_paragraph()
    run = p.add_run(f'[{number:04d}] ')
    run.bold = True
    p.add_run(text)
    return p

def create_patent_document():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    title = doc.add_heading('', 0)
    title_run = title.add_run('发明专利申请书')
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    doc.add_heading('发明名称', level=1)
    doc.add_paragraph('基于离线队列的移动端数据增量同步方法及系统')

    doc.add_heading('技术领域', level=1)
    add_paragraph_with_number(doc, 1,
        '本发明涉及分布式数据同步和移动应用技术领域，尤其涉及一种基于离线队列的移动端数据增量同步方法及系统。')

    doc.add_heading('背景技术', level=1)
    add_paragraph_with_number(doc, 2,
        '移动应用需要在离线和在线状态之间无缝切换，确保用户数据在多设备间保持一致。然而，现有的移动端数据同步技术存在以下问题：')

    add_paragraph_with_number(doc, 3,
        '第一，离线体验差。大多数应用在网络断开时功能严重受限甚至完全不可用，影响用户在地铁、电梯、飞机等弱网环境下的使用体验。')

    add_paragraph_with_number(doc, 4,
        '第二，同步冲突处理不完善。当用户在多个设备上同时修改同一数据，或离线期间的修改与云端数据产生冲突时，现有技术的冲突解决策略简单粗暴，可能导致数据丢失。')

    add_paragraph_with_number(doc, 5,
        '第三，全量同步效率低下。每次同步都传输全部数据，在数据量大或网络带宽有限时，同步过程耗时长、流量消耗大。')

    add_paragraph_with_number(doc, 6,
        '第四，同步失败恢复困难。同步过程中如果网络中断，往往需要重新开始整个同步过程，无法从断点恢复。')

    add_paragraph_with_number(doc, 7,
        '第五，实时性与省电的矛盾。为了及时同步数据，需要频繁唤醒网络连接，导致电池消耗过快。')

    add_paragraph_with_number(doc, 8,
        '因此，需要一种离线优先的增量同步方法，能够在离线时正常工作，在线时高效增量同步，并妥善处理数据冲突。')

    doc.add_heading('发明内容', level=1)
    add_paragraph_with_number(doc, 9,
        '本发明要解决的技术问题是：如何实现移动应用的离线优先架构，确保离线时功能正常，在线时高效增量同步，并妥善处理多设备数据冲突。')

    add_paragraph_with_number(doc, 10,
        '为解决上述技术问题，本发明提供一种基于离线队列的移动端数据增量同步方法，包括以下步骤：')

    add_paragraph_with_number(doc, 11, '步骤S1，离线操作队列管理：')

    add_paragraph_with_number(doc, 12,
        'S1.1 操作队列数据结构：定义离线操作记录结构OfflineOperation，包含字段：id（操作唯一标识，UUID格式）、type（操作类型：CREATE创建/UPDATE更新/DELETE删除）、entity（实体类型：交易/分类/预算/账户）、entityId（实体ID）、payload（操作数据，JSON格式）、timestamp（操作时间戳，精确到毫秒）、retryCount（已重试次数）、status（状态：PENDING待同步/SYNCING同步中/SYNCED已同步/FAILED失败）。')

    add_paragraph_with_number(doc, 13,
        'S1.2 操作入队流程：当用户执行写操作时：（1）立即将数据写入本地SQLite数据库，确保本地数据即时可用；（2）创建对应的OfflineOperation记录，状态设为PENDING；（3）将操作记录插入离线队列（按timestamp排序的优先队列）；（4）如果当前处于在线状态，触发同步服务。')

    add_paragraph_with_number(doc, 14,
        'S1.3 操作合并优化：对同一实体的连续操作进行合并优化。例如：连续两次UPDATE同一交易，合并为最后一次UPDATE；CREATE后紧跟UPDATE，合并为带完整数据的CREATE；CREATE后紧跟DELETE，两个操作互相抵消，从队列中移除。')

    add_paragraph_with_number(doc, 15, '步骤S2，增量变更追踪机制：')

    add_paragraph_with_number(doc, 16,
        'S2.1 版本号机制：每个数据实体维护version字段（从1开始递增）和modifiedAt字段（最后修改时间）。每次修改操作version自增1，modifiedAt更新为当前时间。')

    add_paragraph_with_number(doc, 17,
        'S2.2 客户端同步游标：客户端维护lastSyncTime（上次成功同步的时间戳）和lastSyncVersion（上次同步的全局版本号）。同步时仅请求游标之后的变更数据。')

    add_paragraph_with_number(doc, 18,
        'S2.3 服务端变更日志：服务端维护变更日志表，记录所有数据变更。日志字段包括：changeId、entityType、entityId、changeType、changeData、changeTime、sourceDeviceId。变更日志支持按时间范围高效查询。')

    add_paragraph_with_number(doc, 19, '步骤S3，三阶段同步流程：')

    add_paragraph_with_number(doc, 20,
        'S3.1 上传阶段（Push）：将本地离线队列中状态为PENDING的操作按顺序发送到服务端。每个操作发送后等待服务端确认，成功则标记为SYNCED，失败则根据错误类型决定重试或标记为FAILED。采用批量发送优化，每批最多50个操作。')

    add_paragraph_with_number(doc, 21,
        'S3.2 下载阶段（Pull）：向服务端请求lastSyncTime之后的所有变更。服务端返回变更列表，按changeTime排序。客户端依次应用变更到本地数据库，跳过由本设备产生的变更（通过sourceDeviceId判断）。')

    add_paragraph_with_number(doc, 22,
        'S3.3 确认阶段（Confirm）：所有变更应用成功后，更新lastSyncTime为最新变更的时间戳。清理离线队列中状态为SYNCED的记录。记录本次同步的统计信息（上传条数、下载条数、耗时）。')

    add_paragraph_with_number(doc, 23, '步骤S4，冲突检测与解决：')

    add_paragraph_with_number(doc, 24,
        'S4.1 冲突检测：当同一实体在客户端和服务端都有修改时（即客户端待上传的操作对应的实体，其服务端version高于客户端记录的baseVersion），标记为冲突。')

    add_paragraph_with_number(doc, 25,
        'S4.2 冲突解决策略一——后写入者获胜（默认）：比较客户端操作时间戳和服务端修改时间戳，保留时间较晚的版本。适用于大多数场景，简单高效。')

    add_paragraph_with_number(doc, 26,
        'S4.3 冲突解决策略二——字段级合并：将实体拆分为独立字段，非冲突字段自动合并，仅冲突字段采用后写入者获胜策略。适用于复杂实体，减少数据丢失。')

    add_paragraph_with_number(doc, 27,
        'S4.4 冲突解决策略三——用户决策：将冲突详情展示给用户，包括本地版本、服务端版本、冲突字段高亮，让用户选择保留哪个版本或手动合并。适用于重要数据。')

    add_paragraph_with_number(doc, 28,
        'S4.5 冲突日志记录：无论采用何种策略，都记录冲突解决日志，包括：冲突实体、双方版本、采用的策略、最终结果，便于问题追溯和数据恢复。')

    add_paragraph_with_number(doc, 29, '步骤S5，失败恢复与重试机制：')

    add_paragraph_with_number(doc, 30,
        'S5.1 指数退避重试：操作同步失败后，采用指数退避策略重试。重试间隔序列为：1秒、2秒、4秒、8秒、16秒，最多重试5次。超过重试次数的操作标记为FAILED。')

    add_paragraph_with_number(doc, 31,
        'S5.2 断点续传：同步过程记录进度（当前正在处理的操作ID）。如果同步中断，恢复后从断点继续，已成功的操作不重复处理。支持批量操作的部分成功（每个操作独立确认状态）。')

    add_paragraph_with_number(doc, 32,
        'S5.3 失败操作处理：FAILED状态的操作保留在本地，不影响后续操作的同步。用户可查看失败操作列表，选择：手动重试、放弃该操作、导出失败数据。')

    add_paragraph_with_number(doc, 33, '步骤S6，智能同步触发策略：')

    add_paragraph_with_number(doc, 34,
        'S6.1 网络状态感知：监听网络连接状态变化，从离线切换到在线时自动触发同步。区分WiFi和移动网络，可配置仅WiFi时同步大数据。')

    add_paragraph_with_number(doc, 35,
        'S6.2 电量感知：低电量模式（<20%）下降低同步频率，仅同步关键操作。充电状态下可执行完整同步和数据清理。')

    add_paragraph_with_number(doc, 36,
        'S6.3 批量聚合：短时间内的多次写操作聚合为一次同步请求。聚合窗口默认为3秒，减少网络请求次数。')

    doc.add_heading('附图说明', level=1)
    add_paragraph_with_number(doc, 37, '图1是本发明实施例提供的离线优先架构示意图；')
    add_paragraph_with_number(doc, 38, '图2是本发明实施例提供的三阶段同步流程图；')
    add_paragraph_with_number(doc, 39, '图3是本发明实施例提供的冲突检测与解决流程图；')
    add_paragraph_with_number(doc, 40, '图4是本发明实施例提供的智能同步触发策略图。')

    doc.add_heading('具体实施方式', level=1)
    add_paragraph_with_number(doc, 41, '实施例一：离线记账与自动同步')

    add_paragraph_with_number(doc, 42,
        '用户张先生在地铁（无网络）中使用记账应用记录早餐消费15元：')

    add_paragraph_with_number(doc, 43,
        '本地处理：系统将交易记录写入本地SQLite数据库，生成交易ID=T001。创建OfflineOperation{id:OP001, type:CREATE, entity:Transaction, entityId:T001, payload:{amount:15,category:餐饮,description:早餐}, status:PENDING}。界面立即显示这笔交易，用户体验与在线状态无异。')

    add_paragraph_with_number(doc, 44,
        '地铁到站，手机恢复4G网络。系统检测到网络状态变化，触发同步服务。上传阶段：发送OP001到服务端，服务端返回成功，OP001状态更新为SYNCED。下载阶段：请求lastSyncTime之后的变更，发现妻子在家用平板记录了一笔"水果50元"，合并到本地。确认阶段：更新lastSyncTime，清理已同步的OP001。')

    add_paragraph_with_number(doc, 45, '实施例二：多设备冲突解决')

    add_paragraph_with_number(doc, 46,
        '用户在手机和平板上同时编辑同一笔交易"午餐"，手机将金额改为35元，平板将备注改为"和同事聚餐"：')

    add_paragraph_with_number(doc, 47,
        '手机先完成同步，服务端交易version变为2。平板尝试同步时，检测到冲突（本地baseVersion=1，服务端version=2）。')

    add_paragraph_with_number(doc, 48,
        '系统采用字段级合并策略：金额字段仅手机修改，采用手机版本35元；备注字段仅平板修改，采用平板版本"和同事聚餐"。合并结果：金额35元，备注"和同事聚餐"，version更新为3。')

    add_paragraph_with_number(doc, 49,
        '记录冲突日志：{entityId:T002, conflictFields:[amount,description], resolution:FIELD_MERGE, result:{amount:35,description:和同事聚餐}}。')

    doc.add_heading('有益效果', level=1)
    add_paragraph_with_number(doc, 50, '本发明相比现有技术具有以下有益效果：')
    add_paragraph_with_number(doc, 51, '1. 离线可用：用户在任何网络环境下都可正常使用应用，写操作即时生效，不受网络限制。')
    add_paragraph_with_number(doc, 52, '2. 高效同步：增量同步机制仅传输变更数据，相比全量同步减少90%以上的数据传输量。')
    add_paragraph_with_number(doc, 53, '3. 冲突可控：多策略冲突解决机制，根据数据重要性灵活选择，最大限度保护用户数据。')
    add_paragraph_with_number(doc, 54, '4. 可靠恢复：断点续传和指数退避重试确保同步最终成功，失败操作可追溯处理。')
    add_paragraph_with_number(doc, 55, '5. 省电省流量：智能触发策略平衡同步及时性和资源消耗。')

    # 权利要求书
    doc.add_page_break()
    title = doc.add_heading('', 0)
    title_run = title.add_run('权利要求书')
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run('1. ').bold = True
    p.add_run('一种基于离线队列的移动端数据增量同步方法，其特征在于，包括以下步骤：')
    for item in [
        'a) 构建离线操作队列，用户执行写操作时立即写入本地数据库并创建操作记录入队；',
        'b) 采用版本号和修改时间进行变更追踪，客户端维护同步游标，服务端维护变更日志；',
        'c) 执行三阶段同步流程：上传本地操作、下载服务端变更、确认并更新同步游标；',
        'd) 检测数据冲突并采用多策略解决机制处理冲突；',
        'e) 实现指数退避重试和断点续传的失败恢复机制。'
    ]:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Cm(0.5)

    p = doc.add_paragraph()
    p.add_run('2. ').bold = True
    p.add_run('根据权利要求1所述的方法，其特征在于，所述离线操作队列采用先进先出原则，支持同一实体连续操作的合并优化。')

    p = doc.add_paragraph()
    p.add_run('3. ').bold = True
    p.add_run('根据权利要求1所述的方法，其特征在于，所述增量同步通过比较客户端同步游标与服务端变更日志时间戳实现，仅传输游标之后的变更数据。')

    p = doc.add_paragraph()
    p.add_run('4. ').bold = True
    p.add_run('根据权利要求1所述的方法，其特征在于，所述冲突解决支持三种策略：后写入者获胜、字段级合并和用户决策，并记录冲突解决日志。')

    p = doc.add_paragraph()
    p.add_run('5. ').bold = True
    p.add_run('根据权利要求1所述的方法，其特征在于，所述指数退避重试采用递增间隔序列，所述断点续传记录同步进度支持中断恢复。')

    p = doc.add_paragraph()
    p.add_run('6. ').bold = True
    p.add_run('根据权利要求1所述的方法，其特征在于，还包括智能同步触发策略，根据网络状态、电量状态和操作频率动态调整同步时机。')

    p = doc.add_paragraph()
    p.add_run('7. ').bold = True
    p.add_run('一种基于离线队列的移动端数据增量同步系统，其特征在于，包括：')
    for item in [
        '- 本地存储模块，用于管理本地数据库和离线操作队列；',
        '- 变更追踪模块，用于维护版本号、同步游标和变更日志；',
        '- 同步执行模块，用于执行上传、下载、确认三阶段同步流程；',
        '- 冲突处理模块，用于检测冲突并执行相应解决策略；',
        '- 失败恢复模块，用于实现重试机制和断点续传；',
        '- 触发控制模块，用于根据环境状态智能触发同步。'
    ]:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Cm(0.5)

    p = doc.add_paragraph()
    p.add_run('8. ').bold = True
    p.add_run('一种计算机可读存储介质，其上存储有计算机程序，其特征在于，所述计算机程序被处理器执行时实现权利要求1至6中任一项所述方法的步骤。')

    # 说明书摘要
    doc.add_page_break()
    title = doc.add_heading('', 0)
    title_run = title.add_run('说明书摘要')
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run('本发明公开了一种基于离线队列的移动端数据增量同步方法及系统，属于分布式数据同步技术领域。该方法构建离线操作队列，用户写操作即时写入本地并入队；采用版本号和同步游标实现增量变更追踪；执行上传-下载-确认三阶段同步流程；提供后写入者获胜、字段级合并、用户决策三种冲突解决策略；实现指数退避重试和断点续传的失败恢复机制；根据网络、电量状态智能触发同步。本发明解决了现有技术离线体验差、同步冲突处理不完善、全量同步效率低、失败恢复困难等问题，实现了离线优先、高效可靠的数据同步。')

    doc.save('D:/code/ai-bookkeeping/docs/patents/专利11_离线优先增量同步方法.docx')
    print('专利11文档已生成')

if __name__ == '__main__':
    create_patent_document()
