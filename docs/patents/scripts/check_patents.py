# -*- coding: utf-8 -*-
"""
专利检查脚本 - 根据专利提交检查清单检查专利06-12
"""

import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
import os
import re

def check_patent(filepath):
    """检查单个专利文档"""
    doc = Document(filepath)

    # 收集所有文本
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    text = '\n'.join(full_text)

    results = {
        'filename': os.path.basename(filepath),
        'issues': [],
        'stats': {}
    }

    # ===== 一、形式检查 =====

    # 1. 发明名称（标题）- 不超过25字
    title = ''
    for para in doc.paragraphs:
        if para.style.name.startswith('Heading') or '一种' in para.text[:20]:
            title = para.text.strip()
            break

    if title:
        title_len = len(title.replace(' ', ''))
        results['stats']['title'] = title
        results['stats']['title_length'] = title_len
        if title_len > 25:
            results['issues'].append(f'[形式] 发明名称过长: {title_len}字 (应≤25字)')
    else:
        results['issues'].append('[形式] 未找到发明名称')

    # 2. 摘要字数 150-300字
    abstract_match = re.search(r'说明书摘要(.+?)(?:摘要附图|$)', text, re.DOTALL)
    if abstract_match:
        abstract = abstract_match.group(1).strip()
        # 去除段落编号
        abstract_clean = re.sub(r'\[\d{4}\]\s*', '', abstract)
        abstract_len = len(abstract_clean.replace(' ', '').replace('\n', ''))
        results['stats']['abstract_length'] = abstract_len
        if abstract_len < 150:
            results['issues'].append(f'[形式] 摘要过短: {abstract_len}字 (应150-300字)')
        elif abstract_len > 300:
            results['issues'].append(f'[形式] 摘要过长: {abstract_len}字 (应150-300字)')
    else:
        results['issues'].append('[形式] 未找到说明书摘要')

    # 3. 权利要求格式检查
    claims_section = re.search(r'权利要求书(.+?)(?:说明书摘要|$)', text, re.DOTALL)
    if claims_section:
        claims_text = claims_section.group(1)
        # 统计独立权利要求数
        independent_claims = len(re.findall(r'\n\d+\.\s*一种', claims_text))
        dependent_claims = len(re.findall(r'根据权利要求\d+所述', claims_text))
        total_claims = len(re.findall(r'\n\d+\.', claims_text))

        results['stats']['total_claims'] = total_claims
        results['stats']['independent_claims'] = independent_claims
        results['stats']['dependent_claims'] = dependent_claims

        if independent_claims < 2:
            results['issues'].append(f'[权利要求] 独立权利要求不足: {independent_claims}个 (建议≥2个,含方法和系统)')

        # 检查是否有存储介质权利要求
        if '存储介质' not in claims_text:
            results['issues'].append('[权利要求] 缺少存储介质权利要求')

        # 检查是否有系统权利要求
        if '系统' not in claims_text and '装置' not in claims_text:
            results['issues'].append('[权利要求] 缺少系统/装置权利要求')
    else:
        results['issues'].append('[形式] 未找到权利要求书')

    # 4. 段落编号连续性检查
    para_nums = re.findall(r'\[(\d{4})\]', text)
    if para_nums:
        nums = [int(n) for n in para_nums]
        expected = list(range(1, max(nums)+1))
        actual = sorted(set(nums))
        if actual != expected:
            missing = set(expected) - set(actual)
            if missing:
                results['issues'].append(f'[形式] 段落编号不连续,缺少: {sorted(missing)[:5]}...')
        results['stats']['max_para_num'] = max(nums)
    else:
        results['issues'].append('[形式] 未使用段落编号[0001]格式')

    # 5. 附图标记检查
    figures = re.findall(r'图(\d+)', text)
    if figures:
        fig_nums = sorted(set([int(f) for f in figures]))
        results['stats']['figure_count'] = max(fig_nums)
        expected_figs = list(range(1, max(fig_nums)+1))
        if fig_nums != expected_figs:
            results['issues'].append(f'[形式] 附图编号不连续: {fig_nums}')
    else:
        results['issues'].append('[形式] 未找到附图引用')

    # ===== 二、内容检查 =====

    # 1. 技术领域
    if '技术领域' not in text:
        results['issues'].append('[内容] 缺少技术领域章节')

    # 2. 背景技术 - 检查是否引用专利
    bg_section = re.search(r'背景技术(.+?)发明内容', text, re.DOTALL)
    if bg_section:
        bg_text = bg_section.group(1)
        cn_patents = re.findall(r'CN\d{9}[A-Z]?', bg_text)
        us_patents = re.findall(r'US\d{7,}[A-Z]?\d*', bg_text)
        papers = re.findall(r'论文|paper|学术|IEEE|ACM', bg_text, re.I)

        results['stats']['cn_patents_cited'] = len(cn_patents)
        results['stats']['us_patents_cited'] = len(us_patents)
        results['stats']['papers_cited'] = len(papers)

        if len(cn_patents) + len(us_patents) < 2:
            results['issues'].append(f'[内容] 背景技术专利引用不足: CN{len(cn_patents)}个, US{len(us_patents)}个 (建议≥2个)')
    else:
        results['issues'].append('[内容] 缺少背景技术章节')

    # 3. 技术问题
    if '技术问题' not in text and '共性问题' not in text and '不足' not in text:
        results['issues'].append('[内容] 未明确陈述技术问题')

    # 4. 有益效果 - 检查量化数据
    effects_section = re.search(r'有益效果(.+?)(?:附图说明|权利要求)', text, re.DOTALL)
    if effects_section:
        effects_text = effects_section.group(1)
        percentages = re.findall(r'\d+%', effects_text)
        numbers = re.findall(r'\d+倍|\d+个|\d+秒|\d+ms', effects_text)
        results['stats']['quantified_effects'] = len(percentages) + len(numbers)

        if len(percentages) + len(numbers) < 3:
            results['issues'].append(f'[内容] 有益效果量化数据不足: {len(percentages)+len(numbers)}个 (建议≥3个)')
    else:
        results['issues'].append('[内容] 缺少有益效果章节')

    # 5. 实施例数量
    examples = re.findall(r'实施例\s*\d+', text)
    unique_examples = len(set(examples))
    results['stats']['examples_count'] = unique_examples

    if unique_examples < 3:
        results['issues'].append(f'[内容] 实施例数量不足: {unique_examples}个 (应≥3个)')

    # ===== 三、附图检查 =====

    # 检查附图说明章节
    if '附图说明' not in text:
        results['issues'].append('[附图] 缺少附图说明章节')

    # 检查流程图标记 S101等
    step_marks = re.findall(r'S\d+', text)
    if not step_marks and '步骤' in text:
        results['issues'].append('[附图] 建议使用S101、S102等标准步骤标记')

    return results


def generate_report(results_list):
    """生成检查报告"""
    print('=' * 60)
    print('专利提交检查报告')
    print('=' * 60)

    for result in results_list:
        print(f"\n【{result['filename']}】")
        print('-' * 50)

        # 统计信息
        stats = result['stats']
        if 'title' in stats:
            print(f"  标题: {stats['title']}")
            print(f"  标题长度: {stats.get('title_length', 'N/A')}字")
        print(f"  摘要长度: {stats.get('abstract_length', 'N/A')}字")
        print(f"  权利要求: 共{stats.get('total_claims', 0)}条 "
              f"(独立{stats.get('independent_claims', 0)}条, 从属{stats.get('dependent_claims', 0)}条)")
        print(f"  段落编号: 最大[{stats.get('max_para_num', 0):04d}]")
        print(f"  附图数量: {stats.get('figure_count', 0)}个")
        print(f"  实施例: {stats.get('examples_count', 0)}个")
        print(f"  引用专利: CN{stats.get('cn_patents_cited', 0)}个, US{stats.get('us_patents_cited', 0)}个")
        print(f"  量化效果: {stats.get('quantified_effects', 0)}个")

        # 问题列表
        issues = result['issues']
        if issues:
            print(f"\n  发现问题 ({len(issues)}个):")
            for issue in issues:
                print(f"    ⚠ {issue}")
        else:
            print("\n  ✓ 未发现问题")

    print('\n' + '=' * 60)
    print('检查完成')
    print('=' * 60)


if __name__ == '__main__':
    patent_dir = os.path.dirname(os.path.dirname(__file__))

    # 检查专利06-12
    results = []
    for i in range(6, 13):
        pattern = f'专利{i:02d}_*_完整提交版.docx'
        for f in os.listdir(patent_dir):
            if f.startswith(f'专利{i:02d}_') and f.endswith('_完整提交版.docx'):
                filepath = os.path.join(patent_dir, f)
                print(f'正在检查: {f}')
                result = check_patent(filepath)
                results.append(result)
                break

    print('\n')
    generate_report(results)
