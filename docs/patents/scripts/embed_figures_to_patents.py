# -*- coding: utf-8 -*-
"""
将生成的附图嵌入到专利文档中
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

PATENTS_DIR = 'D:/code/ai-bookkeeping/docs/patents'
FIGURES_DIR = 'D:/code/ai-bookkeeping/docs/patents/figures'

# 专利与附图的映射关系
PATENT_FIGURES = {
    '02': [
        ('专利02_图1_多模态融合系统架构.png', '图1 多模态融合系统架构示意图'),
        ('专利02_图2_多模态融合算法流程.png', '图2 多模态融合算法流程图'),
        ('专利02_图3_置信度融合示意.png', '图3 置信度加权融合示意图'),
    ],
    '03': [
        ('专利03_图1_分层学习架构.png', '图1 分层学习系统架构图'),
        ('专利03_图2_联邦学习流程.png', '图2 联邦学习隐私保护流程图'),
        ('专利03_图3_知识蒸馏过程.png', '图3 知识蒸馏过程示意图'),
    ],
    '04': [
        ('专利04_图1_零基预算架构.png', '图1 零基预算系统架构图'),
        ('专利04_图2_动态分配流程.png', '图2 零基预算分配算法流程图'),
        ('专利04_图3_预算调拨示意.png', '图3 预算跨小金库调拨示意图'),
    ],
    '05': [
        ('专利05_图1_四维交互架构.png', '图1 四维语音交互架构图'),
        ('专利05_图2_语音处理流程.png', '图2 语音处理流程图'),
        ('专利05_图3_对话状态机.png', '图3 多轮对话状态机'),
    ],
    '06': [
        ('专利06_图1_位置增强架构.png', '图1 位置增强财务管理系统架构图'),
        ('专利06_图2_POI匹配算法.png', '图2 POI匹配算法流程图'),
    ],
    '07': [
        ('专利07_图1_去重系统架构.png', '图1 多因子交易去重系统架构图'),
        ('专利07_图2_相似度计算.png', '图2 多因子相似度计算示意图'),
    ],
    '08': [
        ('专利08_图1_可视化架构.png', '图1 财务数据可视化组件架构图'),
        ('专利08_图2_交互手势规范.png', '图2 交互手势规范示意图'),
    ],
    '09': [
        ('专利09_图1_披露层级结构.png', '图1 渐进式披露界面层级结构'),
        ('专利09_图2_行为驱动适配.png', '图2 用户行为驱动的界面适配'),
    ],
    '10': [
        ('专利10_图1_账单解析架构.png', '图1 智能账单解析系统架构图'),
        ('专利10_图2_字段映射算法.png', '图2 字段智能映射算法示意图'),
    ],
    '11': [
        ('专利11_图1_离线同步架构.png', '图1 离线优先增量同步架构图'),
        ('专利11_图2_冲突解决流程.png', '图2 同步冲突解决流程图'),
    ],
    '12': [
        ('专利12_图1_隐私保护架构.png', '图1 隐私保护协同学习架构图'),
        ('专利12_图2_差分隐私机制.png', '图2 差分隐私机制示意图'),
    ],
}

# 专利文件名映射
PATENT_FILES = {
    '02': '专利02_多模态融合智能记账_完整提交版.docx',
    '03': '专利03_分层自学习协同学习_完整提交版.docx',
    '04': '专利04_零基预算动态分配_完整提交版.docx',
    '05': '专利05_四维语音交互_完整提交版.docx',
    '06': '专利06_位置增强财务管理_完整提交版.docx',
    '07': '专利07_多因子交易去重_完整提交版.docx',
    '08': '专利08_财务数据可视化交互_完整提交版.docx',
    '09': '专利09_渐进式披露界面设计_完整提交版.docx',
    '10': '专利10_智能账单解析导入_完整提交版.docx',
    '11': '专利11_离线优先增量同步_完整提交版.docx',
    '12': '专利12_隐私保护协同学习_完整提交版.docx',
}


def add_figures_section(doc, patent_num):
    """在文档末尾添加附图说明部分"""
    figures = PATENT_FIGURES.get(patent_num, [])
    if not figures:
        return

    # 添加附图说明标题
    doc.add_page_break()
    heading = doc.add_heading('附图说明', level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # 添加附图说明文字
    for i, (filename, caption) in enumerate(figures, 1):
        para = doc.add_paragraph()
        para.add_run(f'{caption}：').bold = True
        para.add_run(f'展示了本发明实施例{i}中相关技术方案的具体实现结构和流程。')

    # 添加各附图
    doc.add_page_break()
    heading = doc.add_heading('附图', level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for filename, caption in figures:
        filepath = os.path.join(FIGURES_DIR, filename)
        if os.path.exists(filepath):
            # 添加图片
            doc.add_picture(filepath, width=Inches(6))
            # 居中对齐
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 添加图片标题
            caption_para = doc.add_paragraph()
            caption_run = caption_para.add_run(caption)
            caption_run.bold = True
            caption_run.font.size = Pt(10)
            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 添加空行
            doc.add_paragraph()
        else:
            print(f'  警告: 找不到附图 {filename}')


def process_patent(patent_num):
    """处理单个专利文档"""
    filename = PATENT_FILES.get(patent_num)
    if not filename:
        print(f'  跳过: 专利{patent_num} 无文件映射')
        return

    filepath = os.path.join(PATENTS_DIR, filename)
    if not os.path.exists(filepath):
        print(f'  跳过: 文件不存在 {filename}')
        return

    print(f'  处理: {filename}')

    try:
        # 打开文档
        doc = Document(filepath)

        # 添加附图部分
        add_figures_section(doc, patent_num)

        # 保存带附图的版本
        output_filename = filename.replace('_完整提交版.docx', '_含附图版.docx')
        output_path = os.path.join(PATENTS_DIR, output_filename)
        doc.save(output_path)
        print(f'  已保存: {output_filename}')

    except Exception as e:
        print(f'  错误: {e}')


def main():
    print('='*60)
    print('开始将附图嵌入专利文档...')
    print('='*60)

    for patent_num in PATENT_FIGURES.keys():
        print(f'\n专利{patent_num}:')
        process_patent(patent_num)

    print('\n' + '='*60)
    print('所有专利文档附图嵌入完成！')
    print('='*60)


if __name__ == '__main__':
    main()
